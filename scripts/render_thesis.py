"""Render the graduation thesis to a Word document.

Reproducible single source of truth for the thesis prose. Run:

    uv run python scripts/render_thesis.py \
        --out ~/Desktop/大语言模型驱动智能体的行为模拟研究——以Polymarket预测市场为例.docx

Organisation follows the thesis structure requested by the author:
摘要 / 引言 / 国内外文献综述 / 研究设计 / 实证分析 / 总结与讨论.
The body deliberately contains no engineering-process vocabulary; every
claim is phrased as a research idea, a datum, an analysis, or a conclusion.

All quantitative claims trace to artifacts under docs/v14/ and output/v14/.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = Path(__file__).resolve().parent.parent / "docs" / "v14" / "figures"
TBL_DIR = Path(__file__).resolve().parent.parent / "docs" / "v14" / "tables"


def _load_csv_rows(name, fmt=None):
    """Read a CSV from docs/v14/tables/ and return rows ready for
    three_line_table; optionally format numeric cells via `fmt` dict
    mapping column-name → format string."""
    import csv
    p = TBL_DIR / name
    rows = []
    with p.open(encoding="utf-8") as f:
        for row in csv.DictReader(f):
            out = []
            for k, v in row.items():
                if fmt and k in fmt:
                    try:
                        v = fmt[k] % float(v)
                    except (TypeError, ValueError):
                        pass
                out.append(v)
            rows.append(out)
    return rows


def _load_action_mix_rows():
    return _load_csv_rows("table6_action_mix.csv",
        fmt={a: "%.1f" for a in ["限价单", "市价单", "撤单", "不操作",
                                  "拆分", "合并", "声明信念"]})


def _load_b1_markets_rows():
    rows = _load_csv_rows("table7_b1_markets.csv")
    # shorten the slug for display
    short = {
        "dogecoin-above-0pt34-on-january-17": "狗狗币>0.34",
        "nba-por-mia-2025-01-21": "NBA POR-MIA",
        "will-yamand-orsi-win-the-2024-uruguay-presidential-election": "乌拉圭大选",
        "nba-gsw-min-2025-01-15": "NBA GSW-MIN",
        "will-cducsu-and-spd-form-the-next-german-government": "德国组阁",
        "will-xrp-dip-to-1pt50-in-august-343-253-666-332-591": "XRP<1.50",
        "nba-phx-ind-2025-01-04": "NBA PHX-IND",
        "will-the-price-of-bitcoin-be-less-than-78000-on-mar-28": "BTC<78000",
        "will-xrp-dip-to-1pt00-in-march": "XRP<1.00",
        "will-elon-musk-buy-msnbc-before-april-2025": "马斯克买MSNBC",
    }
    for r in rows:
        r[0] = short.get(r[0], r[0][:14])
    return rows


def _load_archetype_pnl_rows():
    return _load_csv_rows("table8_b3_archetype_pnl.csv")


def _load_scale_rows():
    """table9_scale.csv → display rows for the scale experiment."""
    rows = []
    for r in _load_csv_rows("table9_scale.csv"):
        n, em, es, vol, notion, fills, cancel, spread = r
        rows.append([
            f"{int(float(n))}",
            f"{float(em):.3f} ± {float(es):.3f}",
            f"{float(vol):.4f}",
            f"{float(notion):,.0f}",
            f"{float(fills):.0f}",
            f"{float(cancel):.1f}",
        ])
    return rows


def _load_tick_rows():
    """table10_tick.csv → display rows for the tick-horizon experiment."""
    rows = []
    for r in _load_csv_rows("table10_tick.csv"):
        t, em, es, vol, acts, fills, spread = r
        rows.append([
            f"{int(float(t))}",
            f"{float(em):.3f} ± {float(es):.3f}",
            f"{float(vol):.4f}",
            f"{float(acts):.0f}",
            f"{float(fills):.0f}",
        ])
    return rows


def _load_open_rows():
    """table11_open.csv → display rows for the open-market preview."""
    rows = []
    for r in _load_csv_rows("table11_open.csv"):
        seed, start, end, drift, vol, fills, pnl = r
        rows.append([
            f"种子 {int(float(seed))}",
            f"{float(start):.3f}",
            f"{float(end):.3f}",
            f"{float(drift):+.3f}",
            f"{float(vol):.4f}",
            f"{int(float(fills))}",
        ])
    return rows


# Result paragraph for the profile-distribution experiment (§6.4 五).
# Finalized from the c4 suite once it completes.
_PROFILE_MIX_RESULT = (
    "结果显示,三种画像分布下的仿真终态价格分别为 0.205、0.193、0.247"
    "(均值),彼此之差落在随机性基线之内,价格水平不随画像分布系统性"
    "变化;价格也都停留在初始价附近,没有哪一种分布使价格更接近真实"
    "结果。交易量与成交活跃度在三种分布间的差异同样未超出种子波动。"
    "也就是说,虽然把高赔率追逐型的占比从不足一成提高到七成,显著"
    "改变了群体的画像构成,但市场层面的价格与交易结构几乎不变。这与"
    "本节(一)“用画像群体还是随机群体不产生可检出差异”的结论一致,"
    "并把它推进了一步:不仅画像与随机之分不影响市场结果,画像群体"
    "内部的分布构成也不影响——市场层面结果对参与者的画像构成稳健。"
)


def _load_profile_mix_rows():
    """table12_profile_mix.csv → display rows for the profile-distribution
    experiment."""
    rows = []
    for r in _load_csv_rows("table12_profile_mix.csv"):
        variant, em, es, vol, notion, fills, longshot, spread = r
        rows.append([
            variant,
            f"{float(em):.3f} ± {float(es):.3f}",
            f"{float(vol):.4f}",
            f"{float(notion):,.0f}",
            f"{float(fills):.0f}",
            f"{float(longshot):.1f}",
        ])
    return rows


# Result paragraph for the thinking-mode experiment (§6.4 六).
# Finalized from the c5 suite once it completes.
_THINKING_RESULT = (
    "结果显示,思考模式开与关下的仿真终态价格分别为 0.205、0.205"
    "(均值),差异落在随机性基线之内;价格波动、交易量与动作分布"
    "在两种模式间的差异同样未超出种子波动。也就是说,关闭语言模型"
    "的显式思考步骤,并未改变市场层面的价格与交易结构。这一结果的"
    "意义在于:本工具复现出的交易者行为,主要由结构化的提示、信念"
    "机制与撮合规则决定,而非依赖语言模型在每一轮额外生成一段推理"
    "文字;思考模式更多影响单次决策的可解释性,而非聚合的市场结果。"
)


def _load_thinking_rows():
    """table13_thinking.csv → display rows for the thinking-mode
    experiment."""
    rows = []
    for r in _load_csv_rows("table13_thinking.csv"):
        mode, em, es, vol, notion, fills, spread = r
        rows.append([
            mode,
            f"{float(em):.3f} ± {float(es):.3f}",
            f"{float(vol):.4f}",
            f"{float(notion):,.0f}",
            f"{float(fills):.0f}",
        ])
    return rows


# ----------------------------------------------------------------------
# helpers
# ----------------------------------------------------------------------


# --- NJU 2026 thesis format spec (docs/v13/THESIS_FORMAT_SPEC.md) ---
SONG = "宋体"          # 正文中文
HEI = "黑体"           # 标题中文
KAI = "楷体_GB2312"     # 摘要中文 / 封面信息行
LATIN = "Times New Roman"  # 正文西文与数字
PT_BODY = 12           # 小四
PT_H1 = 16             # 三号(章/参考文献/致谢/附录)
PT_H2 = 14             # 四号(节)
PT_ABS_TITLE = 18      # 摘要页标题
IND_FIRST = Cm(0.85)   # 正文首行缩进 ≈ 2 字符
IND_HANG = Cm(0.76)    # 参考文献悬挂缩进


def _set_run(run, *, ea=SONG, latin=LATIN, size=PT_BODY, bold=False):
    """Apply Chinese + Western font, size, weight to a run, including
    the w:eastAsia attribute python-docx does not set by default."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), ea)


def _page_setup(doc):
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)
        s.header_distance = Cm(1.5)
        s.footer_distance = Cm(1.75)


def h1(doc, text):
    """一级标题。章/参考文献/致谢/附录:黑体三号居中;
    摘要页标题:楷体_GB2312 18pt 居中。"""
    is_abs = ("摘要" in text) or (text.strip().lower() == "abstract")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    if is_abs:
        _set_run(r, ea=KAI, size=PT_ABS_TITLE, bold=True)
    else:
        _set_run(r, ea=HEI, size=PT_H1, bold=True)
    return p


def h2(doc, text):
    """二级标题:黑体四号居左,段前约 0.28 cm。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Cm(0.28)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run(r, ea=HEI, size=PT_H2, bold=True)
    return p


def para(doc, text, indent=True):
    """正文:宋体/Times New Roman 小四,1.5 倍行距,两端对齐,
    首行缩进 ≈2 字符。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = IND_FIRST
    pf.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _set_run(r, ea=SONG, size=PT_BODY)
    return p


def apara(doc, text, lead=None):
    """摘要正文:楷体 12pt,1.5 倍行距。`lead` 为加粗引导词
    (如“摘要：”“关键词：”)。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r0 = p.add_run(lead)
        _set_run(r0, ea=KAI, size=PT_BODY, bold=True)
    r = p.add_run(text)
    _set_run(r, ea=KAI, size=PT_BODY)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    _set_run(r, ea=SONG, size=PT_BODY)
    return p


def _set_cell_border(cell, **edges):
    """Set individual cell borders. edges like top={'sz':12,'val':'single'}."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders")
        tcPr.append(tcB)
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        el = tcB.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcB.append(el)
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(spec.get("sz", 8)))
            el.set(qn("w:color"), spec.get("color", "000000"))


# table number is auto-incremented; caption goes ABOVE (academic style)
_TBL = {"n": 0}
_FIG = {"n": 0}


def three_line_table(doc, title, headers, rows, widths_cm=None):
    """Academic three-line (三线表) table. Caption above; only the
    table top rule, header-bottom rule, and table bottom rule are
    drawn — no vertical or inner horizontal lines."""
    _TBL["n"] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"表 {_TBL['n']}  {title}")
    _set_run(cr, ea=SONG, size=10.5, bold=True)

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = 1  # center
    thick = {"sz": 14, "val": "single"}
    thin = {"sz": 6, "val": "single"}
    n_rows = 1 + len(rows)
    for ci, htext in enumerate(headers):
        c = t.rows[0].cells[ci]
        c.text = str(htext)
        for para_ in c.paragraphs:
            para_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_.paragraph_format.line_spacing = 1.25
            for rr in para_.runs:
                _set_run(rr, ea=SONG, size=10, bold=True)
        _set_cell_border(c, top=thick, bottom=thin)
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(v)
            for para_ in c.paragraphs:
                para_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_.paragraph_format.line_spacing = 1.25
                for rr in para_.runs:
                    _set_run(rr, ea=SONG, size=10)
            bottom = thick if ri == n_rows - 1 else None
            _set_cell_border(c, bottom=bottom)
    if widths_cm:
        for r in t.rows:
            for ci, w in enumerate(widths_cm):
                r.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t


def figure(doc, png_name, title, width_cm=13.5):
    """Embed a figure centered with a Chinese caption BELOW it."""
    _FIG["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / png_name), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"图 {_FIG['n']}  {title}")
    _set_run(cr, ea=SONG, size=10.5, bold=True)
    doc.add_paragraph()


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)


def pagebreak(doc):
    doc.add_page_break()


REFERENCES = [
    "[1] Park J S, O'Brien J C, Cai C J, et al. Generative Agents: Interactive Simulacra of Human Behavior. UIST, 2023.",
    "[2] Hong S, Zhuge M, Chen J, et al. MetaGPT: Meta Programming for a Multi-Agent Collaborative Framework. ICLR, 2024.",
    "[3] Li G, Hammoud H A K, Itani H, et al. CAMEL: Communicative Agents for 'Mind' Exploration of Large Language Model Society. NeurIPS, 2023.",
    "[4] Zhou X, Zhu H, Mathur L, et al. SOTOPIA: Interactive Evaluation for Social Intelligence in Language Agents. ICLR, 2024.",
    "[5] Li N, Gao C, Li M, et al. EconAgent: Large Language Model-Empowered Agents for Simulating Macroeconomic Activities. ACL, 2024.",
    "[6] Piatti G, Jin Z, Kleiman-Weiner M, et al. Cooperate or Collapse: Emergence of Sustainable Cooperation in a Society of LLM Agents. NeurIPS, 2024.",
    "[7] Park J S, Zou C Q, Kamphorst J, et al. LLM Agents Grounded in Self-Reports Enable General-Purpose Simulation of Individuals. arXiv:2411.10109, 2024.",
    "[8] Altera AL, Ahn A, Becker N, et al. Project Sid: Many-agent simulations toward AI civilization. arXiv:2411.00114, 2024.",
    "[9] Piao J, Yan Y, Zhang J, et al. AgentSociety: Large-Scale Simulation of LLM-Driven Generative Agents. arXiv:2502.08691, 2025.",
    "[10] Ashery A F, Aiello L M, Baronchelli A. Emergent social conventions and collective bias in LLM populations. Science Advances, 2025, 11(18).",
    "[11] Vallinder A, Hughes E. Cultural Evolution of Cooperation among LLM Agents. arXiv:2412.10270, 2024.",
    "[12] Fan W, Zhang S, Wang X, et al. AIvilization v0: Toward Large-Scale Artificial Social Simulation. arXiv:2602.10429, 2026.",
    "[13] Yang Z, Zhang Z, Zheng Z, et al. OASIS: Open Agent Social Interaction Simulations with One Million Agents. arXiv:2411.11581, 2024.",
    "[14] Zhang X, Lin J, Mou X, et al. SocioVerse: A World Model for Social Simulation Powered by LLM Agents. arXiv:2504.10157, 2025.",
    "[15] Vezhnevets A S, Agapiou J P, Aharon A, et al. 使用 Concordia 构建行动锚定于物理、社会或数字空间的生成式智能体仿真. arXiv:2312.03664, 2023.",
    "[16] Golechha S, Garriga-Alonso A. Among Us: A Sandbox for Measuring and Detecting Agentic Deception. NeurIPS, 2025.",
    "[17] Zheng S, Trott A, Srinivasa S, et al. The AI Economist: Taxation policy design via two-level deep multiagent reinforcement learning. Science Advances, 2022, 8(41).",
    "[18] Hullman J, Broska D, Sun H, et al. This human study did not involve human subjects: Validating LLM simulations as behavioral evidence. arXiv:2602.15785, 2026.",
    "[19] Anthis J R, Liu R, Richardson S M, et al. LLM Social Simulations Are a Promising Research Method. ICML, 2025.",
    "[20] Guo T, Chen X, Wang Y, et al. Large Language Model based Multi-Agents: A Survey of Progress and Challenges. IJCAI, 2024.",
    "[21] AlKhamissi B, ElNokrashy M, AlKhamissi M, et al. Investigating Cultural Alignment of Large Language Models. arXiv:2402.13231, 2024.",
    "[22] Tan Q, Jiang L, Zeng Y, et al. Mitigating Cultural Bias in LLMs via Multi-Agent Cultural Debate. arXiv:2601.12091, 2026.",
    "[23] Peter S, Riemer K, West J D. The benefits and dangers of anthropomorphic conversational agents. PNAS, 2025, 122(5).",
    "[24] Gao C, Lan X, Li N, et al. 大语言模型赋能的智能体仿真:综述与展望. Humanities and Social Sciences Communications, 2024, 11(1).",
    "[25] Yang Y, Zhang Y, Wu M, et al. TwinMarket: A Scalable Behavioral and Social Simulation for Financial Markets. NeurIPS, 2025. arXiv:2502.01506.",
    "[26] Polymarket. Polymarket 101. https://docs.polymarket.com/cn/polymarket-101, accessed 2026-05.",
    "[27] Polymarket. Markets & Events. https://docs.polymarket.com/cn/concepts/markets-events, accessed 2026-05.",
    "[28] Polymarket. Prices & Orderbook. https://docs.polymarket.com/cn/concepts/prices-orderbook, accessed 2026-05.",
    "[29] Polymarket. Positions & Tokens. https://docs.polymarket.com/concepts/positions-tokens, accessed 2026-05.",
    "[30] Polymarket. Polymarket USD. https://docs.polymarket.com/cn/concepts/pusd, accessed 2026-05.",
    "[31] Polymarket. Order Lifecycle. https://docs.polymarket.com/concepts/order-lifecycle, accessed 2026-05.",
    "[32] Polymarket. Resolution. https://docs.polymarket.com/cn/concepts/resolution, accessed 2026-05.",
]


# ----------------------------------------------------------------------
# document body
# ----------------------------------------------------------------------


def _cover_line(doc, label, value, label_w=4.5, value_w=8.5):
    """A 封面 info line: label aligned right, value left, both 16pt 楷体_GB2312."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label:>4}".replace(" ", "　") + "   ")
    _set_run(r1, ea=KAI, size=16, bold=True)
    r2 = p.add_run(str(value))
    _set_run(r2, ea=KAI, size=16, bold=False)


def cover_page(doc):
    """南京大学本科毕业论文封面。"""
    # 顶部空白
    for _ in range(3):
        doc.add_paragraph()
    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("本 科 毕 业 论 文")
    _set_run(r, ea=SONG, size=26, bold=True)
    # 信息行
    _cover_line(doc, "学    院", "信息管理学院")
    _cover_line(doc, "专    业", "信息管理与信息系统")
    _cover_line(doc, "题    目",
                "大语言模型驱动智能体的行为模拟研究——以 Polymarket 预测市场为例")
    _cover_line(doc, "年    级", "2022     学    号   221820328")
    _cover_line(doc, "学生姓名", "王 友 林")
    _cover_line(doc, "指导教师", "颜 嘉 麒     职    称   教  授")
    _cover_line(doc, "提交日期", "2026 年 6 月")
    pagebreak(doc)


def build(doc: Document) -> None:
    # ---------- 封面 ----------
    cover_page(doc)
    # ---------- 摘要 ----------
    h1(doc, "中文摘要")
    apara(doc,
        "去中心化预测市场通过真实资金交易汇聚参与者对未来事件的判断。"
        "价格、成交、撤单和持仓变化不仅反映市场对结果概率的估计,"
        "也记录了信息和分歧进入市场的过程。以 Polymarket 为代表的"
        "平台公开保留了市场、订单簿、交易和钱包行为等数据,为细粒度"
        "观察预测市场运行提供了条件。与此同时,大语言模型驱动智能体"
        "具备理解事件描述、结合历史状态并生成行动理由的能力,但在金融"
        "交易场景中,这种能力只有转化为受资金、持仓、报价粒度和撮合"
        "规则约束的可执行订单,才具有研究意义。因此,本文关注的问题"
        "不是大语言模型能否直接预测事件结果,而是其驱动的智能体能否"
        "在真实市场规则约束下形成可信的交易行为和市场过程。",
        lead="摘要：")
    apara(doc,
        "本文以 Polymarket 二元预测市场为对象,构建大语言模型驱动的"
        "群体交易模拟系统。系统整合市场目录、历史交易和订单簿等公开"
        "数据,校准初始价格、费率、报价粒度和盘口深度;基于事件发生前的"
        "钱包历史构造参与者画像;将智能体行为限定为限价单、市价单、撤单、"
        "份额拆分与合并、持仓保持和信念更新等可执行操作。围绕交易过程"
        "可信性、智能体数量、决策轮数、关键模块设计以及活跃市场结算前"
        "情景五个问题,本文在已结算市场和未结算市场上开展对照实验,"
        "并以随机种子差异作为噪声基线。")
    apara(doc,
        "研究发现,大语言模型驱动智能体能够在订单簿约束下形成完整、"
        "可记录、可审计的交易过程,但其主要价值在于模拟交易行为和市场"
        "互动,而不是稳定预测最终结算方向。智能体数量增加或决策轮数"
        "延长,主要影响成交活跃度、路径依赖和个体损益分化,并不必然"
        "改善价格方向判断;显式信念机制和初始主观判断生成方式对行为"
        "连续性和价格偏移的影响更为明显。活跃市场实验能够呈现结算前"
        "的价格区间、群体分歧和交易压力,但不应被解释为预测准确率。"
        "本文的实证意义在于说明,智能体交易模拟可以作为观察预测市场"
        "行为机制的补充证据;工具意义在于提供一个可复现、可干预、"
        "可审计的实验环境,用于比较不同市场设置下交易者行为、信息聚合"
        "和市场微观结构的变化。")
    apara(doc, "大语言模型;智能体;去中心化预测市场;Polymarket;"
                "交易行为;市场微观结构",
          lead="关键词：")

    pagebreak(doc)
    h1(doc, "Abstract")
    apara(doc,
        "Decentralised prediction markets aggregate dispersed beliefs through "
        "funded trading: prices, fills, and positions jointly record how "
        "information enters the market. Large-language-model (LLM) agents have "
        "shown promise in social and economic simulations, yet financial "
        "applications require executable orders and compliance with cash, "
        "inventory, and order-book rules. Whether agent populations can "
        "reproduce credible market processes under reproducible conditions "
        "remains an open question that calls for real-data calibration.",
        lead="Abstract: ")
    apara(doc,
        "This thesis adopts a multi-agent simulation approach. Rather than treating "
        "the LLM as a one-shot forecaster, it starts from micro-level trader "
        "behaviour and traces how individual decisions accumulate into prices "
        "and trades in a continuous double-auction book on Polymarket. The "
        "study builds and evaluates an LLM-driven simulation system calibrated "
        "with public market, trade, and order-book data, participant profiles derived "
        "from pre-event wallet histories, and a restricted action set (limit "
        "and market orders, cancellations, split/merge, hold, and belief "
        "updates). Experiments address five questions: credibility "
        "of the trading process, effects of agent count (10/20/50/100), effects "
        "of decision rounds (10/20/50/100), contributions of belief tools, "
        "prior sampling, profile mix, and reasoning mode, and pre-settlement "
        "scenarios on an active market.")
    apara(doc,
        "The findings show that LLM-driven agents can produce complete, "
        "auditable trading processes under order-book constraints, but their "
        "primary value lies in behavioural simulation rather than stable "
        "settlement forecasting; whether the terminal price ends on the "
        "winning side must be distinguished from whether the simulation "
        "path moves toward the winner. Raising agent count or extending "
        "decision rounds mainly affects trading activity and path dependence, "
        "while explicit belief tools and prior sampling have the clearest "
        "effects on behavioural quality. Pre-settlement runs on active markets "
        "can illustrate price ranges and group disagreement before resolution, "
        "but should not be read as forecast accuracy. The system therefore "
        "offers a controllable, reproducible experimental reference for studying "
        "trader behaviour, information aggregation, and market microstructure "
        "in prediction markets, not a substitute for live-market settlement "
        "judgment.")
    apara(doc,
        "large language model; multi-agent simulation; decentralised prediction "
        "market; Polymarket; trading behaviour; market microstructure",
        lead="Keywords: ")

    pagebreak(doc)
    h1(doc, "目录")
    three_line_table(doc, "目录",
        ["章节", "主要内容"],
        [
            ["摘要 / Abstract", "研究目的、方法、主要结果和意义"],
            ["第一章 引言", "研究背景;研究问题;研究方法;研究创新点"],
            ["第二章 国内外文献综述", "大语言模型驱动智能体、金融市场模拟、预测市场与模拟有效性"],
            ["第三章 研究设计", "数据采集;仿真环境;智能体初始化;市场选择;实验设计"],
            ["第四章 实证分析",
             "结束市场复现;智能体规模;模拟轮数;关键模块;活跃市场预演;决策链分析"],
            ["第五章 总结与讨论", "主要研究结论;研究局限;未来研究方向"],
            ["参考文献", "引用文献"],
            ["致谢", ""],
            ["附录 A", "市场用户行为画像的数据定义"],
            ["附录 B", "实验复现说明"],
        ],
        widths_cm=[4.2, 9.8])

    # ====================================================================
    # 第一章 引言
    # ====================================================================
    pagebreak(doc)
    h1(doc, "第一章  引言")

    h2(doc, "一、研究背景")
    para(doc,
        "预测市场是一种以价格聚合分散信息的市场机制。参与者围绕某一"
        "未来事件买卖结果合约,合约价格在一定程度上反映市场对该事件"
        "发生概率的共识判断。与问卷调查、专家访谈或静态模型不同,"
        "预测市场具有真实资金激励和连续交易过程,参与者在信息、风险"
        "偏好、资金规模和交易策略上的差异会通过订单提交、成交、撤单"
        "和持仓变化不断进入市场价格。因此,预测市场不仅可以被视为概率"
        "判断的汇聚机制,也可以被视为观察交易者行为、信息传播和市场"
        "微观结构的实验场。")
    para(doc,
        "Polymarket 是近年来较具代表性的去中心化预测市场之一。平台"
        "以二元结果合约为主要交易对象,使用连续双向拍卖订单簿撮合"
        "交易,并在事件结束后依据外部事实进行结算。相较传统中心化"
        "预测市场,去中心化预测市场的数据具有较强的可追踪性,市场"
        "目录、价格历史、交易记录、钱包行为和持有人信息均可通过公开"
        "接口或链上数据观察。这使得研究者能够在较细粒度上重建市场"
        "运行过程,并将个体钱包的历史行为转化为参与者画像。")
    para(doc,
        "平台用户并不是与平台本身对赌,而是在开放的点对点市场中与其他"
        "用户交易结果份额;份额价格反映市场对事件结果发生概率的集体判断。"
        "Polymarket 上的市场是最基本的可交易单元,通常对应一个“是/否”"
        "二元问题;事件则是组织容器,可以包含一个或多个相关市场。本文"
        "实验以二元市场为单位运行,市场路径标识用于定位具体市场,条件"
        "合约标识和结果代币标识用于关联订单簿、价格和交易数据。")
    three_line_table(doc, "Polymarket 核心概念与本文对应关系",
        ["概念", "平台含义", "本文用途"],
        [
            ["市场", "一个可交易的“是/否”二元问题", "实验运行的基本单位"],
            ["事件", "一个或多个相关市场的组织容器", "理解市场分组"],
            ["市场路径标识", "平台内用于定位市场的文本编码", "筛选市场和组织输出目录"],
            ["条件合约标识", "条件合约中的市场标识", "关联市场、订单簿和成交"],
            ["结果代币标识", "“是/否”结果代币的交易标识", "查询价格、订单簿和成交"],
            ["pUSD", "交易使用的抵押代币", "在仿真中抽象为现金"],
            ["中央限价订单簿", "用户之间提交订单并形成成交的订单簿", "仿真环境的核心机制"],
        ],
        widths_cm=[3.0, 5.2, 5.8])
    para(doc,
        "Polymarket 的价格机制也决定了本文为何将价格路径作为核心观察"
        "对象。平台价格可以解释为概率:每个份额价格位于"
        "0 到 1 之间,“是”份额价格为 0.65 时,可以理解为市场认为事件"
        "发生概率约为 65%。平台显示价格通常取买卖价差的中间价;当"
        "价差较大时,则可能显示最近成交价。因此,本文同时观察买价、"
        "卖价、中间价、价差和成交价,而不只观察最终胜负。")
    para(doc,
        "在交易和资产结构上,Polymarket 使用 pUSD 作为抵押品,并用"
        "结果代币表示“是”和“否”份额。一单位 pUSD 可以拆分生成一组“是”与 "
        "“否”份额;市场结束后,获胜份额可兑换 1 pUSD,失败份额价值为 0。"
        "真实平台中的订单由用户签名后提交到订单簿,运营方在链下撮合,"
        "成交后再通过智能合约在链上结算。本文仿真抽象出现金约束、"
        "结果份额、限价单、市价成交、撤单、份额拆分与合并,以及成交后"
        "的持仓变化。")
    para(doc,
        "Polymarket 的判定机制同样影响本文对实验结果的解释。当事件"
        "结果明确后,市场通过去中心化判定机制确认获胜结果;判定完成后,"
        "获胜代币可以兑换抵押品,失败代币归零。因此,历史已结算市场能够提供明确真实标签,适合用于"
        "检验仿真终态位置和价格移动方向;活跃市场则只能用于结算前"
        "情景分析,不能用于计算预测命中率。")
    para(doc,
        "去中心化市场的另一个重要特征是参与者身份以钱包为基本单位呈现。"
        "Polymarket 采用非托管模式,用户资金由钱包和智能合约控制;实际"
        "使用中,用户可能通过外部账户、代理钱包或智能钱包持有 pUSD 和"
        "结果代币。虽然钱包并不必然等同于自然人,但它保留了参与者"
        "在不同市场中的交易历史、资金规模、价格偏好和持仓行为。对本文"
        "而言,这意味着可以从钱包行为数据出发构造参与者画像,但不能把"
        "画像解释为真实个人的完整心理画像。")
    para(doc,
        "真实预测市场数据能够描述价格、成交和持仓已经如何变化,但仅凭"
        "事后数据较难回答进一步的问题:如果参与者数量不同、交易轮数"
        "不同、信念更新方式不同,市场过程会发生怎样的变化?如果研究者"
        "希望观察交易者行为如何在订单簿中逐步形成价格,仅靠静态预测"
        "或事后统计是不够的,还需要一个能够复现交易规则、控制实验条件"
        "并记录完整过程的模拟环境。")
    para(doc,
        "大语言模型驱动智能体为这一问题提供了新的实现方式。预测市场"
        "中的交易者并非只做数值计算,还会结合事件背景、市场价格、持仓"
        "状态、资金约束和个人判断形成行动。大语言模型可以处理这类"
        "文本与状态混合的信息,但其输出如果只是自然语言判断,仍不能"
        "构成金融交易行为。本文因此不让智能体自由发表观点,而是要求"
        "其在给定资金、持仓和盘口状态下选择结构化交易动作,并由订单簿"
        "环境检验这些动作是否能够执行。")
    para(doc,
        "将大语言模型驱动智能体用于 Polymarket 交易模拟,还需要处理"
        "四个边界问题。第一,智能体决策必须落实为限价单、市价单、撤单、"
        "份额拆分与合并、持仓保持或信念更新等可记录动作。第二,价格"
        "不是单个智能体判断的直接表达,而是多个参与者在多轮交易中共同"
        "作用的结果。第三,参与者画像应尽量来自事件发生前的钱包行为,"
        "避免完全依赖研究者主观设定。第四,模拟不能使用结算后的未来信息,"
        "活跃市场只能用于结算前情景分析,不能提前评价预测命中率。")
    para(doc,
        "基于以上考虑,本文以 Polymarket 为研究对象,构建大语言模型"
        "驱动的预测市场交易模拟系统。本文的目标不是用智能体替代真实"
        "市场作出结算判断,而是在可复现的订单簿环境中观察智能体如何"
        "下单、成交、撤单、调整信念和形成持仓,并进一步分析这些个体"
        "行为如何汇聚为价格路径、成交活跃度和群体分歧。本文关注的是"
        "交易行为能否被可信模拟,以及模拟过程中能够呈现哪些市场机制,"
        "而不是单纯追问能否猜中最终结果。")

    h2(doc, "二、研究问题")
    para(doc,
        "本文所称“大语言模型驱动智能体”,是指以大语言模型为决策核心,"
        "具有参与者画像、市场状态观察、结构化动作空间和跨轮记忆的"
        "模拟交易主体。它不同于一次性文本预测模型,也不同于完全由固定"
        "规则控制的交易程序。本文所称“去中心化预测市场交易行为”,"
        "是指在 Polymarket 二元结果市场中能够被订单簿环境记录和统计"
        "的行为,包括限价挂单、市价成交、撤单、份额拆分与合并、持仓"
        "保持、信念声明与信念更新,以及由这些行为形成的成交、持仓、"
        "损益、价格路径、交易活跃度和资金流关系。本文不把交易行为"
        "简化为预测“是”或“否”,而是将其理解为市场制度约束下连续发生"
        "的行动序列。")
    para(doc,
        "围绕上述界定,本文提出五个研究问题。第一,大语言模型驱动"
        "智能体能否生成可信的去中心化预测市场交易者行为,即其动作分布、"
        "成交行为、持仓变化、损益表现和价格路径能否构成可运转、可解释、"
        "可与真实市场数据对照的市场过程?第二,智能体数量变化会如何影响"
        "成交规模、价格波动、撤单比例、资金流网络和个体损益分布?第三,"
        "在其他条件相同的情况下,决策轮数变化是否会改变价格路径稳定性、"
        "动作结构和结论对停止时点的依赖?第四,显式信念机制、初始主观"
        "判断生成方式、画像分布和显式推理开关等模块,哪些会对行为质量"
        "或市场结果产生超过随机种子噪声的影响?第五,在尚未结算的活跃"
        "市场中,智能体群体能否形成有解释价值的结算前情景,例如价格区间、"
        "群体信念、净买卖压力和交易活跃度?")
    three_line_table(doc, "表 1-1 研究问题与操作化指标",
        ["研究问题", "样本与实验设置", "主要观察指标", "判定要点"],
        [
            ["研究问题一:交易过程是否可信",
             "10 个结束市场,每市场 3 个随机种子(共 30 组)",
             "终态价格、价格移动方向、动作结构、成交、持仓、损益",
             "是否形成可执行、可审计的交易过程"],
            ["研究问题二:智能体规模",
             "Robotaxi 与 Ethereum;智能体数 10/20/50/100",
             "成交名义额、成交笔数、撤单占比、价格波动、损益分化",
             "规模扩大主要增加流动性,还是改变价格结构"],
            ["研究问题三:模拟轮数",
             "Robotaxi 与 Ethereum;决策轮数 10/20/50/100",
             "终态价格、动作总量、逐轮波动、成交笔数",
             "结果是否依赖停止时点"],
            ["研究问题四:关键模块",
             "信念机制、初始判断、画像分布、显式推理开关",
             "撤单、持有不动、信念声明、成交、价格偏移",
             "哪些模块产生超过噪声基线的效应"],
            ["研究问题五:活跃市场情景",
             "Thunder NBA Finals;3 个随机种子",
             "价格区间、漂移方向、成交活跃度、跨种子一致性",
             "能否形成结算前情景证据"],
        ],
        widths_cm=[3.0, 3.6, 4.2, 3.6])
    para(doc,
        "表 1-1 将五个研究问题对应到样本设置、观察指标与判定要点。"
        "本文不以单一预测准确率为唯一评价,而把交易动作、成交、撤单、"
        "损益分化与价格漂移等过程性指标并列纳入,并区分可执行交易过程"
        "与端到端结算预测能力。")

    h2(doc, "三、研究方法")
    para(doc,
        "本文的研究目标在于从微观交易者行为出发,探讨大语言模型驱动智能体"
        "在去中心化预测市场订单簿环境中的行动如何汇聚为价格、成交与"
        "损益等市场结果,并识别影响行为模拟质量的关键设计要素。为实现"
        "该目标,本文在真实 Polymarket 数据校准的基础上搭建多智能体"
        "仿真系统,开展机制对照与要素分析。本文采用以下四种研究方法。")
    para(doc,
        "第一,文献研究法。本文梳理大语言模型智能体、预测市场行为仿真"
        "以及仿真有效性评价等相关研究,归纳现有工作在"
        "可执行交易约束、真实数据校准和推断边界方面的进展与不足,为"
        "仿真框架的设计和评价指标的选择提供理论基础。")
    para(doc,
        "第二,多源数据采集与描述性统计方法。去中心化预测市场的仿真"
        "须以可复查的真实市场状态和参与者历史为基础。本文整合 Polymarket "
        "的市场目录、交易数据与订单簿等公开数据源,构建涵盖市场目录、成交、"
        "价格历史、持有人与订单簿等信息的实证数据库,并据此导出市场"
        "初始参数与钱包行为特征,为仿真环境设定和参与者画像构造提供"
        "校准依据。")
    para(doc,
        "第三,智能体交易模拟方法。该方法通过"
        "刻画异质智能体在制度规则下的互动,分析宏观市场现象背后的微观"
        "生成机制。本文在连续双向拍卖订单簿中,将大语言模型作为智能体"
        "的决策核心,把输出约束为限价单、市价单、撤单、份额拆分与合并、"
        "持仓保持及信念更新等可执行动作,使语言模型的情境理解服从于"
        "交易制度和资金持仓约束,而非停留于一次性文本预测。")
    para(doc,
        "第四,聚类分析与对照实验方法。在参与者层面,本文对事件前钱包"
        "行为特征进行标准化处理,并采用 K-means 聚类划分行为画像;在系统"
        "层面,本文围绕智能体规模、模拟轮数、信念机制、主观判断生成"
        "方式、画像分布和显式推理开关等变量,在多个结束市场与 Robotaxi、"
        "Ethereum 等基底市场上开展单因素对照实验,并以随机种子差异估计"
        "不可约噪声基线,区分可识别的处理效应与残余波动。")
    para(doc,
        "上述四类方法相互衔接:文献研究界定问题与边界,数据采集与描述性"
        "统计提供校准输入,智能体交易模拟生成可观察的市场过程,聚类与对照实验"
        "则分别刻画参与者异质性与机制效应。其具体实现环节构成第三章"
        "所述技术路线,并作为第四章实证分析的依据。总体流程如下。")
    figure(doc, "fig1_loop.png", "图 1-1 仿真技术路线", width_cm=15)
    para(doc,
        "图 1-1 概括数据采集、参数校准、画像构造、智能体初始化、订单簿"
        "仿真与结果分析六个环节。后续章节的价格、成交、动作和损益指标,"
        "均可追溯到图中相应节点。")

    h2(doc, "四、研究创新点")
    para(doc,
        "第一,本文将大语言模型驱动智能体置于去中心化预测市场的真实"
        "交易规则中,并把智能体输出严格约束为可执行的交易动作。已有"
        "大语言模型智能体研究多集中于文本互动、社会模拟和协作任务,"
        "即使涉及经济场景,也较少把智能体放入具有真实订单簿约束的交易"
        "环境。本文将语言模型决策与连续双向拍卖机制结合,使每个智能体"
        "决策都能被撮合、撤销、成交或结算,从而区分“语言模型能否给出"
        "判断”和“智能体能否参与市场过程”。")
    para(doc,
        "第二,本文建立了真实数据校准的预测市场仿真流程。市场参数、"
        "初始盘口、交易轮数、参与者画像和初始主观判断均尽量由真实数据"
        "导出。参与者画像只使用目标事件发生前的数据构造,避免未来"
        "信息泄漏。")
    para(doc,
        "第三,本文不仅观察最终价格,还系统记录和分析动作结构、成交"
        "规模、撤单行为、持仓变化、个体损益、价格路径和资金流网络。"
        "本文发现智能体可以形成可记录、可审计、可分析的交易过程,"
        "但这并不意味着其能够稳定判断最终结算方向。该边界有助于"
        "避免将交易行为模拟误解为直接预测工具。")
    para(doc,
        "第四,本文通过规模实验、时长实验和多组模块实验识别影响模拟"
        "质量的关键机制。显式信念机制能够改善行为连续性,初始主观判断"
        "生成方式会影响价格是否发生系统性偏移,而画像分布、显式推理等"
        "设计的影响则在当前重复次数下仍需谨慎解释。")
    para(doc,
        "第五,本文将活跃市场纳入实验对象,考察结算前情景分析的可行性。"
        "在真实结果尚未出现时,仿真仍可输出价格区间、交易压力和群体信念"
        "等结构化证据;但其结论须待市场结束后再评价准确性,避免将情景"
        "分析误写为预测保证。")

    # ====================================================================
    # 第二章 国内外文献综述
    # ====================================================================
    pagebreak(doc)
    h1(doc, "第二章  国内外文献综述")
    para(doc,
        "预测市场研究关注价格如何聚合分散信息,以及交易者如何在订单簿"
        "中通过报价、成交、撤单和持仓变化表达判断;大语言模型驱动"
        "智能体研究则关注语言模型能否在具备记忆、目标、工具和环境反馈"
        "的条件下持续行动。本文的研究问题位于两类文献的交汇处:如果"
        "将智能体放入真实数据校准的预测市场制度中,它们能否生成可执行、"
        "可审计的交易行为,并使个体判断在多轮互动中汇聚为价格、成交与"
        "损益等市场结果。本章据此不按文献时间顺序罗列,而围绕三个问题"
        "展开:第一,大语言模型智能体何以能够作为持续行动主体;第二,"
        "智能体群体模拟和金融市场模拟如何处理异质行为与制度约束;第三,"
        "模拟结果应如何验证,以及哪些结论不能被解释为预测能力。")
    h2(doc, "一、大语言模型驱动智能体研究")
    para(doc,
        "Park 等提出 Generative Agents,将大语言模型与记忆、反思和计划"
        "机制结合,使智能体在虚拟环境中表现出较连贯的日常行为与社会"
        "互动[1]。MetaGPT、CAMEL 与 SOTOPIA 分别从协作流程、角色扮演"
        "与社会智能评估角度扩展多智能体系统[2][3][4];Guo 等与 Gao 等"
        "的综述指出,行为质量取决于记忆、角色、环境反馈与工具约束"
        "[20][24]。")
    para(doc,
        "EconAgent 以 LLM 赋能主体模拟宏观经济活动[5];AgentSociety、"
        "Project Sid、OASIS、SocioVerse 与 AIvilization 等工作将仿真"
        "规模推向城市级乃至百万智能体[8][9][12][13][14];Concordia "
        "强调行动须锚定于可执行操作空间[15]。这些研究的环境多为宏观"
        "指标或社交平台,而非连续双向拍卖订单簿,参与者通常不受现金、"
        "持仓与最小报价单位约束,因此其结论不能直接外推至预测市场。")
    para(doc,
        "Cooperate or Collapse 等研究考察 LLM 群体中的合作与规范演化"
        "[6][10][11];Park 等基于自报告资料构建个体模拟体[7];Golechha "
        "等讨论智能体欺骗等行为风险[16]。相关结论提醒本文:群体结果"
        "须通过订单簿中的成交与价格观察,且“画像”“信念”应界定为"
        "操作性构造[23]。")
    para(doc,
        "由此可见,大语言模型智能体研究为本文提供的不是“让模型直接"
        "判断事件结果”的依据,而是“让模型在环境中持续行动”的方法基础。"
        "对于预测市场而言,关键并不在于智能体是否能给出一个看似合理的"
        "“是”或“否”判断,而在于它能否根据市场问题、当前价格、持仓状态"
        "和风险约束,选择可执行动作并在后续反馈中调整行为。因此,本文"
        "将智能体输出限制为交易环境能够处理的结构化动作,并把画像、信念"
        "和记忆视为行为生成机制,而非真实心理状态的直接还原。")
    h2(doc, "二、智能体群体模拟、金融市场与预测市场模拟")
    para(doc,
        "智能体群体模拟从异质智能体的局部规则出发,研究微观互动"
        "如何形成宏观模式,适用于真实数据难以支撑反事实干预的情形"
        "[24]。AI Economist 以多智能体强化学习模拟税收制度[17]。"
        "这些研究说明,如果研究对象涉及异质主体、制度约束和连续互动,"
        "仅依赖事后统计往往难以观察机制,需要能够改变局部条件并记录"
        "过程的模拟系统。")
    para(doc,
        "**预测市场与 Polymarket 机制。** 预测市场的特殊性在于,合约价格"
        "既是交易价格,也是对事件发生概率的市场化表达。Polymarket 进一步"
        "将这一机制落实为二元结果份额、中央限价订单簿、钱包持仓、链下"
        "撮合与事件判定等制度安排[26][27][28][29][31][32]。这使其不同于"
        "一般股票或宏观经济模拟:研究者不仅可以观察价格路径,还可以观察"
        "订单簿深度、买卖价差、成交记录、钱包历史和结算结果之间的关系。"
        "对本文而言,这一结构提供了两个重要条件:一是可以用事件前数据"
        "构造参与者画像和市场初始状态,降低未来信息泄漏风险;二是已结算"
        "市场提供明确结果标签,使仿真终态位置、价格移动方向和交易过程"
        "质量能够被区分评价。")
    para(doc,
        "**金融市场行为模拟。** Yang 等提出的 TwinMarket 将 LLM 智能体"
        "用于金融市场行为与社会互动模拟,强调异质交易者互动对价格与成交"
        "的影响,是少数直接面向金融场景的生成式模拟工作[25]。该方向与"
        "本文高度相关,但本文的研究对象进一步限定为去中心化预测市场。"
        "预测市场中的二元合约价格具有概率解释,事件结束后又存在明确"
        "结算结果,因此“价格路径是否呈现过程性发现”与“终态是否位于"
        "真实胜方一侧”可以分开考察。本文在 Polymarket 真实数据上校准"
        "盘口、费率、价格刻度、订单簿深度与钱包画像,并将智能体动作限定"
        "为挂单、成交、撤单、份额拆分与合并、持仓保持和信念更新等"
        "可审计操作,以此回应研究问题一至五所涉及的交易过程可信性、"
        "规模效应、轮数效应、模块设计和活跃市场结算前情景分析。")
    h2(doc, "三、LLM 仿真的验证与边界")
    para(doc,
        "Hullman 等指出,将仿真结果当作行为证据时须审慎处理验证标准"
        "与推断边界[18];Anthis 等强调须建立严谨评价体系[19]。文化"
        "对齐与拟人化风险研究提示,模型输出不能等同于真实主体心理"
        "[21][22][23]。上述讨论与本文的评价方式直接相关。研究问题一"
        "要求检验智能体是否真正进入交易制度,而不是只生成交易理由;"
        "因此本文观察动作结构、订单有效性、成交、撤单、持仓和价格路径。"
        "研究问题二至四关注不同设计变量的影响,因此需要通过对照实验和"
        "随机种子重复区分处理效应与残余波动。研究问题五面对尚未结算"
        "市场,只能讨论结算前价格区间、交易压力和群体分歧,不能计算"
        "预测命中率。由此,本文不把终态价格是否“猜对”作为唯一标准,"
        "而将模拟定位为预测市场行为研究的实验工具:它可以帮助观察制度"
        "约束下的交易过程,但不能替代真实市场作出结算判断。")
    h2(doc, "四、本章小结")
    para(doc,
        "综上,现有研究已经证明大语言模型智能体可以在记忆、角色、工具"
        "和环境反馈支持下参与持续互动,也说明智能体群体模拟适合观察"
        "微观行为如何形成宏观结果;金融市场生成式模拟则进一步提示,"
        "大语言模型智能体可被用于交易者异质行为研究。但对于去中心化"
        "预测市场而言,仍存在三个不足:一是缺少在真实 Polymarket 数据"
        "校准下的订单簿级交易模拟;二是缺少将钱包历史、盘口状态和结算"
        "规则同时纳入的可复现实验设计;三是容易把模拟终态误读为预测"
        "能力,而忽视动作、成交、持仓和价格路径等过程证据。本文正是在"
        "这一空白上展开:以 Polymarket 连续双向拍卖环境为制度约束,"
        "检验大语言模型驱动智能体能否形成可信交易过程,并进一步分析"
        "智能体规模、模拟轮数、信念机制、画像分布和活跃市场情景输出"
        "的含义。下一章将把上述文献脉络落实为数据采集、订单簿环境、"
        "智能体初始化和实验矩阵设计。")

    # ====================================================================
    # 第三章 研究设计
    # ====================================================================
    pagebreak(doc)
    h1(doc, "第三章  研究设计")

    h2(doc, "一、数据采集")
    para(doc,
        "本文的数据采集围绕 Polymarket 市场运行所需的三类信息展开:"
        "市场信息、交易信息和参与者信息。市场信息用于确定仿真对象和"
        "结算结果,交易信息用于导出价格路径、初始盘口和行为对照,"
        "参与者信息用于构造智能体画像和初始化参数。为了避免将未来"
        "信息引入仿真,所有用于智能体初始化的特征均以目标市场开始时间"
        "为截断点,只保留事件发生前可观察的数据。")
    para(doc,
        "这些数据源与 Polymarket 的平台结构一一对应。Gamma API 中的"
        "事件、市场路径标识和结算字段用于确定研究哪一个问题;订单簿相关"
        "数据中的结果代币标识、最小报价单位、买卖盘和中间价用于确定这个问题"
        "如何被交易;Data API 中的钱包、成交和持仓数据用于确定谁在交易"
        "以及如何交易。因此,本文不能只处理网页标题,还必须同时处理市场"
        "标识、结果代币、订单簿和钱包行为。")
    three_line_table(doc, "数据来源、记录规模与研究用途",
        ["数据类别", "记录数", "包含信息", "研究用途"],
        [
            ["市场基础数据",       "146,231",    "市场名称、状态、结算结果", "市场筛选与真值标注"],
            ["市场交易规则数据",   "1,050,851",  "结果份额、价格刻度、费率", "导出可交易参数"],
            ["成交交易数据",       "42,042,912", "成交价、数量、钱包地址",   "真实成交、钱包历史和价格对照"],
            ["价格历史数据",       "16,075,541", "历史中间价与时间序列",     "早期价格锚点和价格路径"],
            ["参与者与持仓数据",   "2,895,288",  "钱包名称、持有人、持仓",   "画像补充与持仓状态"],
            ["订单簿数据",         "2,732,298",  "买卖盘、价差、盘口深度",   "初始流动性状态"],
        ],
        widths_cm=[3.0, 2.2, 4.8, 4.2])
    para(doc,
        "本文使用的数据可以按研究用途概括为六类。成交交易数据规模最大,"
        "达 4204 万条,是真实成交、钱包历史与价格对照的主要来源;价格"
        "历史数据和市场交易规则数据也达到百万级,分别提供价格锚点"
        "和最小报价单位、交易费率等市场参数。各类数据之间通过市场"
        "标识、结果代币和钱包地址关联,从而在“市场—成交—参与者”三层"
        "之间建立可追踪的对应关系。表中数据规模也决定了本文的市场"
        "筛选与画像构造可以在事件前截断条件下取得足够大的可用样本。")
    para(doc,
        "数据处理过程主要包括市场候选池构建、市场参数导出、钱包特征"
        "计算和仿真结果记录四步。市场候选池首先排除结算信息不清、"
        "交易记录过少、起始价格已接近确定结果的市场;随后保留在事件"
        "开始前具有足够钱包历史、价格记录和盘口信息的市场。这样做的"
        "目的,是保证后续模拟既有可校准的市场环境,也有可用于构造"
        "智能体画像的参与者行为基础。对于缺少完整盘口的市场,本文采用"
        "“优先使用真实盘口、缺失时使用成交价格回退”的规则导出初始"
        "参数。所有用于智能体初始化的钱包特征均以目标市场开始时间为"
        "截断点,避免使用事件发生后的信息。")

    h2(doc, "二、仿真环境构建")
    para(doc,
        "本文的仿真环境是一个二元结果市场的连续双向拍卖订单簿。"
        "市场中有“是”和“否”两侧结果份额,事件结算后获胜侧价值为 1,"
        "失败侧价值为 0。智能体可以用现金买入某一侧份额,也可以卖出"
        "已有份额。这里的现金对应真实平台中的 pUSD 抵押品,YES 和 NO "
        "持仓对应真实平台中的结果代币。环境负责维护订单簿、现金、"
        "持仓、未成交订单和成交记录,并按照价格优先、时间优先的规则"
        "撮合买卖委托。")
    para(doc,
        "真实 Polymarket 的中央限价订单簿采用链下撮合、链上结算的混合结构,"
        "订单需要签名、提交、撮合、结算和确认。本文仿真保留其中"
        "与市场行为最相关的逻辑环节,即订单提交、盘口排队、成交撮合、"
        "撤单和持仓更新;签名验证、链上授权、交易哈希和最终确认等工程"
        "环节不进入模型。这样做的原因是,本文研究对象是交易行为和市场"
        "微观过程,而不是区块链交易执行系统本身。")
    three_line_table(doc, "智能体动作设计及其研究作用",
        ["动作类型", "对应的市场行为", "执行约束", "可用于观察的指标"],
        [
            ["限价挂单", "表达愿意成交的价格,并向订单簿提供流动性",
             "必须给出方向、结果份额、价格和数量", "报价深度、挂单方行为、成交状态"],
            ["市价成交", "表达即时交易意愿,观察是否愿意承担价差成本",
             "只能按当前盘口成交,并受现金或持仓限制", "吃单方行为、即时成交、价格冲击"],
            ["撤销挂单", "允许观点变化或风险控制时收回未成交委托",
             "只能撤销自身尚未成交的订单", "撤单比例、挂撤循环、盘口深度变化"],
            ["持有不动", "保留观望选择,避免强迫每轮交易",
             "不改变订单簿和持仓", "持有不动占比、观望倾向"],
            ["份额拆分", "模拟用抵押品生成 YES/NO 配对份额",
             "受现金余额限制", "可交易库存变化"],
            ["份额合并", "模拟将配对份额回收为现金",
             "需要同时持有等量 YES 和 NO", "现金回收、风险暴露下降"],
            ["信念更新", "保留跨轮概率判断,减少每轮重新猜测",
             "需要给出 YES 概率、置信度和简短理由", "信念路径、交易一致性"],
        ],
        widths_cm=[2.4, 4.3, 4.3, 3.6])
    para(doc,
        "上述动作并不是程序接口字段的简单罗列,而是对预测市场中关键"
        "交易行为的研究化抽象。限价挂单、市价成交和撤销挂单对应交易者"
        "提供流动性、消耗流动性和调整风险暴露的行为;份额拆分与份额合并"
        "对应预测市场中抵押品和结果份额之间的转换;持有不动保留了交易者"
        "观望或等待信息的可能;信念更新则把智能体的概率判断显式记录"
        "下来,使研究者能够比较“声明判断、交易方向和成交结果”之间是否"
        "一致。通过这些动作,语言模型的自然语言判断被转化为可执行、"
        "可校验、可统计的市场行为。")
    para(doc,
        "智能体在每一轮决策时收到两类信息:固定信息包括行为画像、"
        "风险偏好、市场问题、结算规则、可用动作、硬性交易约束和基本"
        "策略提示;动态信息包括当前两侧最优买价、最优卖价和中间价、"
        "最近几轮价格序列、市场进度、自身现金和持仓、未成交订单数、"
        "最近若干轮动作摘要,以及已声明的信念状态。智能体无法直接看到"
        "其他智能体身份和内部状态,也不会被告知真实结算结果。")
    para(doc,
        "撮合环境遵循可执行性、可记录性和可复现性三个原则。智能体"
        "输出必须是结构化动作而非自由文本;每个动作、成交、持仓"
        "和价格变化都会写入实验输出;实验配置、随机种子、市场参数和"
        "模型设置均被记录,保证同一配置可以被再次运行和审查。")
    para(doc,
        "每一轮都是一次离散的市场观察和决策窗口。每轮中,所有"
        "智能体依次观察同一时点附近的市场状态并作出决策,环境执行"
        "这些动作后进入下一轮。每一轮结束后,环境会写入本轮动作、"
        "成交、持仓和价格状态,并把最新盘口、最近价格序列、未成交"
        "订单、现金持仓和已声明信念带入下一轮提示。增加轮数并不是"
        "简单地增加自然时间,而是增加智能体观察和行动的机会。")
    three_line_table(doc, "单轮模拟执行流程与记录内容",
        ["执行阶段", "环境处理", "实际示例", "记录结果"],
        [
            ["状态观察", "读取盘口、近期价格、现金、持仓和上一轮信念",
             "YES 中间价 0.575、现金 100、无未成交订单", "形成本轮提示上下文"],
            ["智能体决策", "根据画像、信念和市场状态选择动作",
             "高频广覆盖型智能体选择小额限价买入", "生成动作、方向、价格、数量和理由"],
            ["动作校验", "检查现金、持仓、价格范围和订单归属",
             "买入金额超过现金时被拒绝或转为空操作", "记录有效或无效动作"],
            ["订单撮合", "按价格优先、时间优先撮合订单",
             "市价买入“是”份额吃掉当前最低卖单", "写入挂单方、吃单方、成交价和数量"],
            ["状态回写", "更新现金、持仓、未成交订单和最新价格",
             "买方现金减少、YES 持仓增加", "形成下一轮可见状态"],
            ["结果输出", "写入动作、成交、持仓和价格路径",
             "第 5 轮产生 22 个动作、3 笔成交", "支撑后续统计和图表分析"],
        ],
        widths_cm=[2.4, 4.1, 4.2, 3.8])

    h2(doc, "三、智能体初始化")
    para(doc,
        "智能体初始化是本文模拟系统的核心环节之一。若全部智能体由研究者"
        "手工设定,模拟容易变成主观角色扮演;若直接使用目标市场运行过程"
        "中的交易行为,又会引入未来信息。为此,本文采用事件前钱包历史"
        "构造智能体:先在目标市场开始时间之前截断数据,再从钱包的历史"
        "成交规模、市场覆盖、价格偏好和交易时间分布中提取行为特征,"
        "最后将这些特征转化为资金规模、风险倾向、行为画像和初始主观"
        "判断。这样既保留真实参与者的异质性,又避免让智能体提前知道"
        "目标市场的后续结果。")
    three_line_table(doc, "智能体初始化使用的钱包行为特征",
        ["特征", "含义", "对初始化的作用"],
        [
            ["累计名义交易额",       "钱包历史交易规模",        "影响资金规模和活跃度"],
            ["最大单一市场集中度",   "资金是否集中在少数市场",   "区分集中型与分散型参与者"],
            ["单位资金市场广度",     "单位资金覆盖市场数量",    "刻画市场探索范围"],
            ["平均成交价",           "历史交易价格中心",        "反映高价/低价偏好"],
            ["尾部极端价格占比",     "极端赔率交易比例",        "识别高赔率追逐倾向"],
            ["活跃时长",             "钱包持续交易时间",        "区分长期与短期参与者"],
            ["成交价波动",           "历史成交价分散程度",      "反映策略稳定性"],
            ["burstiness（交易爆发度）", "交易时间集中程度",       "区分爆发型与持续型参与者"],
        ],
        widths_cm=[3.6, 5.0, 5.6])
    para(doc,
        "用于聚类的八个钱包行为特征分别刻画不同交易侧面。前七个"
        "特征覆盖资金规模、集中度、广度、价格倾向、极端价偏好、活跃"
        "时长与价格波动等维度;burstiness 在原有七个特征基础上"
        "补充刻画交易时间分布的爆发程度。引入该指标后,"
        "样本市场上可稳定得到六类画像;在仅使用七个特征时"
        "部分市场上的聚类数会退化为两类。第三列“对初始化"
        "的作用”说明每一个特征会进入智能体的资金规模、活跃度或交易"
        "偏好等初始参数,从而保证仿真中的群体异质性源自真实钱包"
        "行为分布。")
    para(doc,
        "画像构造依次包括事件前截断、特征计算、标准化处理、聚类分组"
        "和文本画像生成。首先,系统以目标市场开始时间为截断点,只保留"
        "该时间之前的钱包交易记录;其次,对每个钱包计算资金规模、市场"
        "广度、价格偏好和交易时间分布等行为特征;再次,对连续特征进行"
        "截尾处理和标准化,降低极端大户对聚类中心的影响;最后使用 "
        "K-means 聚类形成行为相近的钱包群体,并将每一类的统计特征"
        "转写为智能体可读取的简短行为画像。")
    three_line_table(doc, "画像聚类流程与示例",
        ["画像构造环节", "处理内容", "实际示例", "输出结果"],
        [
            ["事件前截断", "只保留目标市场开始前的钱包交易记录",
             "目标市场开始后的交易不进入画像计算", "无未来信息泄漏的钱包样本"],
            ["特征计算", "计算资金规模、市场广度和成交价偏好等特征",
             "历史成交额 200 美元、参与 9 个市场、均价 0.603", "钱包行为特征向量"],
            ["标准化与截尾", "降低极端大户或异常价格对聚类中心的影响",
             "将极端成交额压缩到分位数边界内", "可比较的标准化特征"],
            ["聚类分组", "使用 K-means 将行为相近的钱包分到同一类",
             "广覆盖活跃钱包归入 C1,高价大额配置钱包归入 C5", "六类行为画像"],
            ["文本画像生成", "把统计特征转写为智能体可读描述",
             "偏好较高价格结果合约,资金规模较大", "初始化提示中的行为描述"],
        ],
        widths_cm=[2.6, 4.0, 4.4, 3.4])
    para(doc,
        "在本文样本中,画像构建稳定得到六类行为画像。下表以最新"
        "事件前钱包池的聚类结果为例,列出每类画像的总体占比和核心"
        "统计特征。表中的成交额、市场数、集中度、均价和极端价占比"
        "均为该类钱包的中位数或比例指标,用于说明不同画像之间的"
        "行为差异。为评估画像结构的因果作用,本文还构造了均匀和"
        "集中两种对照分布。")
    three_line_table(doc, "六类钱包行为画像的统计特征",
        ["画像类型", "钱包占比", "中位成交额", "中位参与市场数", "单一市场集中度", "历史平均成交价", "极端价格交易占比"],
        [
            ["C0 高价极端集中型", "13.5%", "70",  "2",  "91.1%", "0.993", "100.0%"],
            ["C1 广覆盖活跃型",   "31.8%", "200", "9",  "35.5%", "0.603", "64.3%"],
            ["C2 中间价常规型",   "18.2%", "47",  "3",  "60.7%", "0.504", "0.0%"],
            ["C3 低价逆向型",     "12.2%", "18",  "3",  "51.2%", "0.007", "100.0%"],
            ["C4 高频广覆盖型",   "4.8%",  "460", "57", "11.0%", "0.556", "33.3%"],
            ["C5 高价大额配置型", "19.6%", "465", "8",  "29.4%", "0.987", "100.0%"],
        ],
        widths_cm=[3.0, 1.4, 1.7, 1.7, 2.0, 1.9, 2.0])
    para(doc,
        "六类画像并非主观命名,而是由真实钱包行为差异支撑。"
        "其中“高价”与“低价”只描述历史成交价格区间,不直接代表"
        "真实交易者心理上的乐观或悲观。"
        "C1 与 C4 的市场覆盖更广,其中 C4 的中位参与市场数达到 57 个;"
        "C0 和 C5 的平均成交价接近 1 且极端价交易占比很高,但 C5 的"
        "中位成交额更大;C3 的平均成交价接近 0,体现出明显低价逆向"
        "特征;C2 则更接近中间价格和低极端价交易的常规型。")
    para(doc,
        "初始主观判断的生成方式是本文在实验中发现的一个关键方法问题。"
        "早期实现使用截断重抽,即先围绕目标均值抽取概率值,若落在"
        "0 和 1 之外则丢弃重取。在均值接近边界时该方法会系统性改变"
        "实际均值,使智能体集体出现非预期方向偏移。本文随后改用"
        "均值保持的有界分布,使抽样结果在 0 到 1 内合法,同时保持"
        "设定的目标均值。本文还引入显式信念机制,智能体可主动执行"
        "信念更新动作,声明当前 YES 概率、置信度和理由,并在后续轮次"
        "中将该信念回填到提示。这一机制减少了反复挂撤和空转,并为"
        "研究者提供了一个可观察的中间变量。")
    para(doc,
        "在提示构造上,本文尽量避免诱导智能体知道真实结果。智能体会"
        "看到市场问题、结算规则、当前价格、近期价格和自身状态,但不会"
        "看到目标市场未来真实结算,也不会看到用未来数据计算出的画像。"
        "它收到的初始主观判断被明确表述为起始看法,而不是事实答案。")
    three_line_table(doc, "提示构成及实际示例",
        ["提示组成", "写入内容", "实际示例", "设计目的"],
        [
            ["市场问题与规则", "市场题目、YES/NO 含义和结算条件",
             "雷霆队是否会赢得 2026 年 NBA 总冠军", "保证智能体围绕同一事件判断"],
            ["当前市场状态", "中间价、最优买卖价、价差和近期价格",
             "YES 中间价 0.575,最近价格小幅波动", "提供可交易环境和价格锚点"],
            ["自身资金与持仓", "现金、YES/NO 持仓、未成交订单",
             "现金 100,YES 持仓 12,NO 持仓 5", "防止决策脱离余额和库存约束"],
            ["行为画像", "画像类别、风险倾向和交易风格",
             "高频广覆盖型,偏好小额多市场参与", "引入来自真实钱包的异质性"],
            ["信念状态", "上轮 YES 概率、置信度和简短理由",
             "上一轮信念 0.58,置信度 0.60", "保持跨轮判断连续性"],
            ["动作约束", "可选动作、价格范围和数量限制",
             "可限价、市价、撤单、拆分、合并、更新信念或持有", "保证输出可被订单簿执行"],
        ],
        widths_cm=[2.6, 4.0, 4.4, 3.4])
    para(doc,
        "本文还将智能体理由控制在简短范围内。原因在于,过长的自然语言"
        "推理会增加成本,也可能使模型沉浸在自我解释中而忽略动作执行。"
        "交易环境更需要明确动作和参数,因此本文要求智能体输出结构化"
        "动作,并附带简短理由。理由用于辅助分析,真正进入市场的是"
        "经过校验的动作。")

    h2(doc, "四、市场选择")
    para(doc,
        "在完成智能体初始化方案后,本文再确定实验市场。这样安排的"
        "原因在于,市场选择不仅取决于题材和结算状态,也取决于该市场"
        "在事件开始前是否能够提供足够的钱包历史来构造智能体群体。"
        "市场选择并非孤立抽样,而是需要同时满足“有可用"
        "市场参数”和“有可用参与者画像”两个条件。")
    para(doc,
        "本文将纳入实证分析的市场分为三层:结束市场,"
        "用于跨市场行为复现;受控机制市场(Robotaxi 与 Ethereum),"
        "作为规模、时长、画像分布和显式推理实验的共同基准;以及"
        "活跃市场(Thunder NBA Finals),用于开展结算前"
        "情景分析。")
    three_line_table(doc, "市场样本分层与功能定位",
        ["市场类型", "代表市场", "选择理由", "论文作用"],
        [
            ["结束", "Tesla FSD、BTC 100k 等",
             "题材与价格区间分散,且有真实结算结果",
             "检验跨市场行为复现"],
            ["受控", "Tesla Robotaxi、Ethereum $5K",
             "起点和钱包样本较稳定,适合重复改变单一变量",
             "作为规模、轮数、画像和显式推理实验基准"],
            ["活跃", "OKC Thunder 2026 NBA Finals",
             "实验时尚未结算,处于真实信息更新过程中",
             "展示结算前情景分析用途"],
        ],
        widths_cm=[2.8, 3.4, 4.5, 3.5])
    three_line_table(doc, "实证分析所用市场清单",
        ["市场类型", "中文题目", "事件题材", "选择原因"],
        [
            ["结束", "特斯拉是否会在 10 月 31 日前推出无人监督 FSD", "科技/企业",
             "高起点市场,检验终态侧与移动方向差异"],
            ["结束", "特朗普是否会在周一前向华盛顿特区部署国民警卫队", "政治",
             "起点接近 0.5,事件叙事清晰"],
            ["结束", "最高法院是否会作出支持特朗普关税案的裁决", "司法",
             "制度性事件,补充政治司法题材"],
            ["结束", "Katy Perry 与 Trudeau 是否在 10 月底前确认关系", "名人",
             "非金融叙事市场,检验迁移性"],
            ["结束", "比特币是否会在 7 月达到 12.5 万美元", "加密资产",
             "价格型问题,交易活跃"],
            ["结束", "比特币到 2025 年底是否会一直高于 10 万美元", "加密资产",
             "同题材不同结算条件"],
            ["结束", "MicroStrategy 是否会在 8 月 26 日至 9 月 1 日购买比特币", "企业/加密",
             "高起点且真实结果为否"],
            ["结束", "NFL 堪萨斯城酋长队对纽约巨人队比赛结果市场", "体育",
             "补充体育题材样本"],
            ["结束", "NFL 巴尔的摩乌鸦队对布法罗比尔队比赛结果市场", "体育",
             "接近中间价,同题材对照"],
            ["结束", "Lord Miles 是否会在沙漠中完成 40 天只喝水禁食", "公共事件",
             "个人行为承诺类事件"],
            ["结束", "Tesla 是否会在 10 月 31 日前推出无人驾驶 Robotaxi", "科技/企业",
             "受控机制实验共同基准"],
            ["结束", "以太坊是否会在 8 月达到 5000 美元", "加密资产",
             "与 Robotaxi 形成题材对照"],
            ["活跃", "OKC 雷霆是否会赢得 2026 年 NBA 总冠军", "体育/活跃",
             "无结算标签,用于结算前情景分析"],
        ],
        widths_cm=[2.4, 6.0, 2.2, 3.6])
    para(doc,
        "这些市场的选择遵循三个原则。第一,题材上尽量分散,覆盖科技、"
        "政治、司法、娱乐、加密资产和体育事件,避免结论只由单一叙事"
        "类型驱动。第二,价格区间上包含高起点、低起点和接近 0.5 的"
        "市场,因为终态价格是否位于正确侧会受到起始价格强烈影响。"
        "第三,实验用途上区分跨市场检验样本、受控机制市场和活跃市场:前者"
        "用于检验不同题材下的行为复现能力,中者用于改变单一变量,后者用于展示"
        "结算前情景分析能力。")

    h2(doc, "五、实验设计")
    para(doc,
        "本文实验围绕前述五个研究问题展开,整体上分为结束市场复现、"
        "规模实验、时长实验、模块消融实验和活跃市场预演五类。所有"
        "实验配置遵循“只改变一个因素”的原则:规模实验只改变智能体"
        "数量,时长实验只改变决策轮数,消融实验只改变对应模块的"
        "开关或方式。")
    three_line_table(doc, "研究问题、样本市场与处理变量对应关系",
        ["实验类型", "样本市场", "处理变量", "随机种子"],
        [
            ["结束市场复现",  "10 个结束市场(多题材样本)", "市场题材", "每市场 3 组"],
            ["研究问题二：智能体规模", "Robotaxi 与 Ethereum", "智能体数 10/20/50/100", "每档 3 组"],
            ["研究问题三：模拟轮数", "Robotaxi 与 Ethereum", "决策轮数 10/20/50/100", "每档 3 组"],
            ["研究问题四：画像分布", "Robotaxi 与 Ethereum", "自然/均匀/集中", "每档 3 组"],
            ["研究问题四：显式推理", "Robotaxi 与 Ethereum", "链式推理开启/关闭", "每档 3 组"],
            ["活跃市场预演",  "Thunder NBA Finals", "—", "3 组"],
        ],
        widths_cm=[3.2, 4.2, 4.0, 2.2])
    para(doc,
        "下表概括各类实验的处理变量与样本范围。研究问题二与研究问题三的智能体"
        "数量和决策轮数各取 10/20/50/100 四档;研究问题四的画像分布取自然、"
        "均匀和集中三种,显式推理取链式推理开启和关闭两档。受控实验均在 "
        "Robotaxi 与 Ethereum 两个基底市场上重复实施,以便比较同一处理在不同"
        "起点价格下的表现。各组仿真均保留动作、成交、持仓与价格路径等完整"
        "记录,供事后复核与对照分析。")
    para(doc,
        "在所有实验之前,本文首先估计随机性基线。在配置完全固定、决策"
        "温度较低的情况下,仅改变随机种子,仿真终态 YES 中间价仍会"
        "出现可观波动。这说明语言模型后端、群体抽样过程、轮内处理"
        "顺序和订单簿撮合路径都会引入不可约噪声。本文将随机种子差异"
        "作为组间比较的解释基线,凡小于该量级或方向不一致的差异均不"
        "作强因果解释。每次仿真至少产生智能体动作记录、成交记录、"
        "持仓记录和初始化画像记录四类结果数据;基于这些输出,本文可以"
        "同时进行微观行为分析和宏观市场结果分析。")

    # ====================================================================
    # 第四章 实证分析
    # ====================================================================
    pagebreak(doc)
    h1(doc, "第四章  实证分析")
    para(doc,
        "本章围绕五个研究问题展开实证分析,但论证重点不是逐一解释图表,"
        "而是说明模拟系统在何种意义上能够模拟预测市场交易过程,以及这种"
        "能力的边界在哪里。为此,本章按照“宏观结果—受控因素—模块机制—"
        "微观决策链”的顺序展开:先利用结束市场检验智能体能否在订单簿"
        "约束下形成完整交易过程,再考察智能体数量和模拟轮数是否改变"
        "价格、成交与路径依赖,随后分析信念机制、初始判断、画像分布和"
        "显式推理等模块的作用,最后通过决策链分析解释个体信念如何转化"
        "为交易行为。")
    para(doc,
        "本章的核心判断是:大语言模型驱动智能体能够生成可执行、可记录、"
        "可审计的交易过程,但交易过程可信并不等同于稳定预测结算方向。"
        "终态价格是否位于胜方一侧,可能在很大程度上继承真实市场起始"
        "价格中的信息;模拟过程中价格是否朝真实结果移动,才更接近过程性"
        "价格发现。因此,本章把终态价格、价格漂移、成交、撤单、持仓、"
        "信念声明和动作序列共同作为证据,而不以单张图或单个指标决定结论。")
    para(doc,
        "依照这一思路,结束市场实验回答“能否形成可信交易过程”,规模和"
        "轮数实验回答“更多参与者或更长运行是否改变市场结果”,模块实验"
        "回答“哪些设计真正影响行为质量”,活跃市场实验回答“系统能否用于"
        "结算前情景分析”,决策链分析则解释“为什么个体行为看似自洽,"
        "群体价格发现仍然有限”。图表在本章中主要承担证据展示和结果核验"
        "作用,正文叙述以研究问题之间的递进关系为主线。")

    h2(doc, "一、结束市场复现:交易过程可信,过程性价格发现有限")
    para(doc,
        "结束市场复现实验在 10 个市场上各运行 3 个随机种子,共 30 组运行。"
        "结果显示,所有运行均产生了完整的动作、成交、持仓和价格记录,"
        "平均每组约 68 笔成交、约 1497 美元名义成交额,动作结构以信念更新和限价挂单为主,"
        "同时包含市价成交、撤单和持仓保持。这说明智能体不是只输出"
        "静态判断,而是在订单簿环境中持续观察、声明信念并提交可执行"
        "动作。由此可以先得到第一层结论:智能体能够在交易制度内持续"
        "行动,而不是只给出一次性判断。下一步需要判断的是,这种完整"
        "过程是否同时意味着有效的价格发现。")
    figure(doc, "行为可信性_跨市场价格路径与真值移动.png",
           "十市场归一化价格路径与朝真值移动占比",
           width_cm=15)
    para(doc,
        "价格路径显示,多数市场的终态价格并未远离起点。10 个市场的"
        "终态均值从 0.233(BTC 12.5 万)到 0.832(MicroStrategy 购币)"
        "不等,但这些终态大多接近各自初始价格。若按市场均值观察,"
        "起点—终点方向朝真值移动的市场有 8 个;若进一步拆分到 30 组"
        "随机种子,只有 13 组朝真值漂移,但有 21 组终态位于胜方一侧。"
        "这说明“终态位于胜方一侧”和“运行过程朝真值移动”必须分开"
        "解释。前者在相当程度上继承了真实市场起点中的信息,后者才更"
        "接近本文所说的过程性价格发现。")
    figure(doc, "行为可信性_十市场起点终点漂移与真值标注.png",
           "十市场起点价与终点价漂移（含真值标注）", width_cm=15)
    para(doc,
        "将起点价与终点价并置后,上述判断更加清楚。多数市场的终点"
        "与起点之差在 ±0.04 以内,只有“Trump 部署国民警卫队”一组"
        "从 0.500 明显下移至 0.237。其余市场基本贴近对角线,说明起点"
        "价格被较强地保留到终点。在 10 个市场、3 个随机种子的样本"
        "规模下,价格发现幅度普遍较小。因此后续规模实验和轮数实验的"
        "核心问题,就是观察增加参与者或延长运行是否能够改变这种起点"
        "锚定。")
    para(doc,
        "从价格方向看,结果需要谨慎解释。按本地结算侧标注,30 组运行"
        "中有 21 组终态价格位于胜方一侧,但只有 13 组价格朝胜方方向"
        "移动。前者说明终态位置经常保留了真实市场起始价格中的信息,"
        "后者说明智能体群体在仿真过程中并未稳定发现结算方向。整体"
        "111 次仿真的起点—终点散点也呈现同一现象:大量市场终态贴近"
        "对角线,说明价格主要在起点附近波动而非主动发现。")
    figure(doc, "行为可信性_全部仿真起点价与终点价关系.png",
           "全部仿真起点价—终点价散点（绿=真值 YES，红=真值 NO）",
           width_cm=12)
    para(doc,
        "把观察范围扩大到全部 111 次仿真后,起点锚定仍然存在。起点—"
        "终点平均绝对差仅 0.052;真值为 YES 的 54 次仿真中,终点高于"
        "0.5 的只有 11 次;真值为 NO 的 57 次仿真中,终点低于 0.5 的"
        "只有 18 次。换言之,仿真终态更常停留在初始价格附近,而不是"
        "被推向胜方对应的边界。这个更大范围的散点证据与前述 30 组"
        "结束市场结果一致,共同说明价格发现强度弱于初始锚定。")
    para(doc,
        "综上,LLM 驱动智能体可以生成可执行、可审计、可分析的交易行为"
        "和市场过程,但不能据此推出其具备稳定的结算方向判断能力。"
        "这一结论构成后续实验的解释前提:规模、轮数和模块设计即使改变"
        "成交活跃度或动作结构,也未必会改变价格发现边界。")
    para(doc,
        "上述结果说明,模拟系统首先通过了“能否形成交易过程”的检验,"
        "但也暴露出价格发现弱于起点锚定的问题。接下来需要进一步考察"
        "的是,这一边界是否只是由于参与者数量不足造成的。如果增加智能体"
        "数量能够显著改变成交密度与价格中心,那么规模可能是改善市场结果"
        "的关键;反之,如果规模只放大交易活跃度而不改变方向,则说明系统"
        "瓶颈不在参与者数量本身。")

    h2(doc, "二、智能体规模:成交活跃度上升,价格方向未见系统性变化")
    para(doc,
        "本节结论:智能体数量增加显著提高成交笔数与名义额,但终态价格"
        "未呈现随规模单调改善方向的规律。研究问题二的规模实验在 "
        "Robotaxi 与 Ethereum 两个市场上固定 20 轮、"
        "画像类型、信念机制、温度和其他环境参数,仅改变智能体数量,"
        "智能体数为 10/20/50/100,每档 3 个种子。结果显示,动作数几乎随"
        "智能体数量线性增长,成交笔数和成交名义额也明显上升;但终态"
        "价格并不呈现“数量越多越接近某个方向”的单调规律。")
    figure(doc, "规模效应_智能体数量对价格成交和盈亏分化的影响.png",
           "各规模下终态价、成交额与盈亏分化",
           width_cm=15)
    para(doc,
        "规模变化首先反映在交易强度上,而不是价格方向上。Robotaxi 的"
        "终态 YES 价格始终位于 0.405–0.450 之间,Ethereum 位于"
        "0.523–0.560 之间,均接近各自起点;与此同时,成交名义额随"
        "智能体数量快速上升,Ethereum 在 100 个智能体时达到 36791 USD,"
        "约为 10 个智能体时的 73 倍。个体损益分散也随规模扩大而增加,"
        "说明更多智能体主要放大交易活动和盈亏尾部,并没有稳定推动"
        "价格中心移动。")
    figure(doc, "规模效应_不同智能体数量下价格路径全景.png",
           "各规模下 YES 中间价轨迹（多种子叠加）",
           width_cm=15)
    para(doc,
        "逐轮价格路径进一步表明,规模扩大没有带来新的收敛方向。"
        "Robotaxi 的轨迹大多被压缩在 0.39–0.51 的窄带内,Ethereum "
        "则位于 0.46–0.61 附近。智能体数量增加后,单条轨迹的波动"
        "幅度略有下降,例如 Robotaxi 的终态标准差从 N=10 的 0.038 "
        "降至 N=100 的 0.009,但价格中心仍然没有随规模发生单调移动。")
    three_line_table(doc, "规模实验核心指标（每档 3 种子，均值±标准差）",
        ["市场", "智能体数", "终态价", "成交名义额", "成交笔数", "撤单比例"],
        [
            ["Robotaxi", "10",  "0.435 ± 0.038", "327",    "17.7",  "7.17"],
            ["Robotaxi", "20",  "0.433 ± 0.033", "1038",   "43.0",  "6.65"],
            ["Robotaxi", "50",  "0.405 ± 0.017", "7042",   "163",   "5.99"],
            ["Robotaxi", "100", "0.450 ± 0.009", "11111",  "363",   "5.81"],
            ["Ethereum", "10",  "0.548 ± 0.019", "503",    "32.3",  "6.85"],
            ["Ethereum", "20",  "0.560 ± 0.036", "2344",   "74.7",  "6.63"],
            ["Ethereum", "50",  "0.538 ± 0.063", "11854",  "172",   "6.37"],
            ["Ethereum", "100", "0.523 ± 0.035", "36791",  "416",   "5.67"],
        ],
        widths_cm=[2.4, 1.2, 3.0, 2.6, 2.0, 2.0])
    para(doc,
        "核心指标也支持这一判断。终态价在两市场上分别局限于 0.405–"
        "0.450 与 0.523–0.560 区间,变化幅度小于随机性基线;成交名义"
        "额与成交笔数则随智能体数量单调上升。Ethereum 从 10 个智能体"
        "增至 100 个智能体时,名义额从 503 USD 增至 36791 USD,成交"
        "笔数从 32.3 增至 416。撤单占比反而略降,说明智能体增多后"
        "挂单后再撤单的循环没有同步放大。")
    figure(doc, "规模效应_不同智能体数量下个体盈亏分布.png",
           "各规模下智能体盈亏分布", width_cm=15)
    para(doc,
        "损益分布揭示了规模效应的另一个侧面:规模越大,尾部越宽,"
        "但中位智能体仍接近零收益。Robotaxi 在 50 个智能体时的损益"
        "极差扩大到 −1289 至 +444,Ethereum 在 100 个智能体时扩大到"
        " −1128 至 +6873。由此可见,规模扩大放大了少数高频或大额"
        "智能体的盈亏贡献,但并未改变多数智能体接近零收益的分布中心。")
    figure(doc, "规模效应_不同智能体数量下撤单率随轮数变化.png",
           "各规模下撤单率随轮次变化", width_cm=15)
    para(doc,
        "撤单行为则呈现出早期较高、随后趋稳的特点。Robotaxi 的平均"
        "撤单率从 N=10 的约 7.29% 降至 N=100 的约 5.74%,Ethereum "
        "也从约 6.73% 降至约 5.63%。这说明规模扩大后,撤单并没有成为"
        "主导行为;更多智能体带来的是更多成交机会和更高交易额,而不是"
        "更混乱的订单撤回。")
    para(doc,
        "规模实验表明,智能体数量主要改变流动性和交互密度,并不自动"
        "提高价格发现质量。更多智能体会带来更多报价机会、更多成交"
        "和更大的名义交易额,但价格方向仍受到个体初始信念、轮内撮合"
        "路径和随机种子的影响。规模扩大放大的是个体损益的离散程度,"
        "而非市场整体的价格收敛。")
    para(doc,
        "规模实验表明,更多智能体主要提高流动性和交互密度,而没有稳定"
        "推动价格方向。由此自然引出第二个受控因素:如果不是“参与者更多”,"
        "那么“互动更久”是否会让价格逐步吸收智能体信念并趋于稳定?模拟"
        "轮数实验正是为了区分成交累积、反馈链条拉长与真实价格发现之间"
        "的关系。")

    h2(doc, "三、模拟轮数:累计成交增加,长时运行可能放大路径依赖")
    para(doc,
        "本节结论:延长决策轮数会增加累计成交,Robotaxi 上终态价格缓慢下移,"
        "Ethereum 长时运行则伴随更明显的种子间分歧。研究问题三的时长实验"
        "固定 20 个智能体和其他环境参数,仅改变决策轮数,"
        "分别为 10/20/50/100,每档 3 个种子。Robotaxi 与 Ethereum "
        "两个市场的四档轮数均已闭合。结果显示更长的市场会带来更多"
        "观察和行动机会,但其结果不是简单收敛。")
    figure(doc, "时长效应_模拟轮数对价格路径和路径依赖的影响.png",
           "各轮数下终态价格与归一化均价路径",
           width_cm=15)
    para(doc,
        "轮数实验显示,更长的运行会增加反馈链条,但并不保证结果更稳定。"
        "Robotaxi 的终态价随轮数由 0.450 降至 0.388,变化幅度为 0.062,"
        "接近随机性基线边缘;Ethereum 的变化更大,终态价从 0.583 降至"
        "0.193。归一化均价路径显示,Ethereum 的 100 轮运行在中后段"
        "持续下行,说明长运行在该市场上放大了累计漂移。两市场都出现"
        "长轮次下移,但幅度差异很大,因此不能把轮数简单理解为价格发现"
        "能力的增强。")
    figure(doc, "时长效应_不同模拟轮数下价格路径全景.png",
           "各轮数下价格轨迹（多种子，归一化进度）", width_cm=15)
    para(doc,
        "种子间差异是解释长轮次结果的关键。Robotaxi 四档轮数的轨迹"
        "宽度大致相当,终态标准差为 0.010–0.048;Ethereum 的 50 轮和"
        "100 轮轨迹宽度明显放大,终态标准差达到 0.240 与 0.283,个别"
        "种子价格甚至下行至接近 0。也就是说,Ethereum 长轮数下的整体"
        "下移并不是所有种子均匀下行,而是若干路径被不断放大后的结果。")
    three_line_table(doc, "时长实验核心指标（每档 3 种子，均值±标准差）",
        ["市场", "决策轮数", "终态价", "成交名义额", "成交笔数"],
        [
            ["Robotaxi", "10",  "0.450 ± 0.010", "629",  "21.0"],
            ["Robotaxi", "20",  "0.440 ± 0.020", "1128", "45.7"],
            ["Robotaxi", "50",  "0.428 ± 0.035", "1880", "93.3"],
            ["Robotaxi", "100", "0.388 ± 0.048", "2689", "131"],
            ["Ethereum", "10",  "0.583 ± 0.034", "1011", "22.7"],
            ["Ethereum", "20",  "0.545 ± 0.043", "1218", "53.3"],
            ["Ethereum", "50",  "0.367 ± 0.240", "2528", "110"],
            ["Ethereum", "100", "0.193 ± 0.283", "5526", "183"],
        ],
        widths_cm=[2.4, 1.4, 3.0, 2.8, 2.4])
    para(doc,
        "数值上,轮数增加稳定推高成交规模。Robotaxi 成交名义额从"
        "629 增至 2689 USD,Ethereum 从 1011 增至 5526 USD;成交笔数"
        "也相应增长,Robotaxi 从 21.0 增至 131,Ethereum 从 22.7 增至"
        "183。与此同时,Ethereum 的终态标准差从 0.034 扩大到 0.283,"
        "说明行动机会增加的同时,路径依赖也被放大。")
    figure(doc, "时长效应_不同模拟轮数下个体盈亏分布.png",
           "各轮数下智能体盈亏分布", width_cm=15)
    para(doc,
        "个体损益同样体现了长轮次的累积效应。Robotaxi 各档损益"
        "中位数接近 0,标准差在 11–30 之间;Ethereum 在 100 轮时损益"
        "标准差扩大到约 741,极差达到 −2421 至 +3781。少数智能体的"
        "极端盈亏与价格路径中的种子差异相互对应,进一步说明长时模拟"
        "并不必然给出更稳定的市场结果。")
    para(doc,
        "Robotaxi 市场中,轮数从 10 增至 100 后,成交笔数和成交额"
        "逐步增加,终态价格也出现更明显下移;但每轮动作量基本稳定,"
        "说明总动作增加主要来自运行轮数拉长。Ethereum 的 50 轮和"
        "100 轮结果均显示更强负向漂移,且种子间差异很大。"
        "时长实验说明,决策轮数会影响累计成交和价格路径,长时运行"
        "可能放大信念修正、撮合路径和随机种子差异,不能简单认为"
        "“更长轮次必然带来稳定结果”。")
    para(doc,
        "轮数实验进一步表明,延长运行会增加成交和反馈机会,但也可能"
        "放大路径依赖。仅靠规模和时长不足以解释行为质量差异,因此需要"
        "回到智能体内部设计,考察信念状态、初始主观判断、画像构成和"
        "显式推理等模块是否会改变行动结构与市场结果。")

    h2(doc, "四、关键模块:信念与先验影响明确,画像与推理模式主要改变动作结构")
    para(doc,
        "本节结论:显式信念机制与初始主观判断生成方式对行为连续性与价格"
        "偏移影响最为清晰;画像分布主要改变活跃度,对终态方向影响不稳定;"
        "显式推理开关主要改变动作结构与成交笔数,不能稳定改善价格发现。"
        "本文将模块结果分为两类:机制证据较清晰的(信念、先验),以及影响"
        "存在但方向不稳的(画像、显式推理)。")
    para(doc,
        "**显式信念机制**的作用不是让智能体预测得更准,而是让它在"
        "多轮交易中保留一个可回看的概率状态,减少每轮重新猜测造成的"
        "行为断裂。前序信念开关实验显示,开启信念机制后撤单和空转"
        "行为下降;显式推理实验亦显示,动作结构会随推理开关"
        "发生系统变化。")
    para(doc,
        "**初始主观判断生成方式**的影响则体现在价格偏移上。早期"
        "截断重抽会在均值靠近 0 或 1 时改变实际抽样均值,把靠近边界"
        "的判断往中间推,从而让部分市场出现非预期买压。改用均值保持"
        "的有界抽样后,价格偏移明显下降。该修正的意义不是让系统直接"
        "知道真实结果,而是防止初始化方法把额外偏差注入智能体群体。")
    para(doc,
        "**画像分布实验**(研究问题四)显示,初始化画像比例会改变参与者构成,"
        "并在部分市场上影响波动率、成交额或终态价,但当前样本不足以"
        "支持稳定价格效应。自然、均匀与集中三种分布之间存在差异,"
        "但缺少跨市场一致方向。")
    figure(doc, "模块消融_画像分布对市场结果的影响.png",
           "三种画像分布下的终态 YES 价格", width_cm=15)
    para(doc,
        "画像构成的变化没有带来稳定的价格方向变化。Robotaxi 上三种"
        "分布的终价均值分别为 0.417、0.450、0.440,跨度 0.033;"
        "Ethereum 上分别为 0.540、0.527、0.585,跨度 0.058。三种"
        "分布之间的差距与每档内的 3 种子标准差量级(0.008–0.041)"
        "基本相当。因此,把画像构成从“自然分布”改成“均匀”或“集中”,"
        "尚不足以在两个基底市场上支持稳定的价格效应。")
    three_line_table(doc, "画像分布实验核心指标",
        ["市场", "画像分布", "完整种子", "终态价", "波动率", "成交额"],
        [
            ["Robotaxi", "自然分布", "3", "0.417 ± 0.024", "0.0143", "946"],
            ["Robotaxi", "均匀分布", "3", "0.450 ± 0.015", "0.0254", "1708"],
            ["Robotaxi", "集中分布", "3", "0.440 ± 0.041", "0.0129", "2045"],
            ["Ethereum", "自然分布", "3", "0.540 ± 0.035", "0.0083", "1400"],
            ["Ethereum", "均匀分布", "3", "0.527 ± 0.008", "0.0100", "5570"],
            ["Ethereum", "集中分布", "3", "0.585 ± 0.025", "0.0087", "1866"],
        ],
        widths_cm=[2.0, 2.8, 1.8, 3.0, 2.0, 2.0])
    para(doc,
        "成交额的变化比终态价格更明显。Robotaxi 的均匀分布与集中"
        "分布成交额(1708、2045 USD)明显高于自然分布(946 USD),"
        "Ethereum 均匀分布上的成交额(5570 USD)也显著高于另两种"
        "(1400、1866 USD)。这说明画像构成确实会改变交易活跃度,"
        "但价格中心位置基本不变。由此,画像分布更适合被解释为"
        "群体构成与活跃度因素,而不是决定终态方向的因素。")
    para(doc,
        "**显式推理实验**(研究问题四)的结果更清楚地体现在动作结构上。"
        "开启链式推理后,两个市场中的持有不动占比上升,信念更新"
        "占比下降,撤单比例也低于关闭组;同时成交笔数增加。但终态"
        "价格方向并不一致,因此不能说显式推理改善了价格发现。")
    figure(doc, "模块消融_思考模式对动作结构的影响.png",
           "显式推理开启与关闭下的动作结构对比", width_cm=15)
    para(doc,
        "显式推理开关首先改变的是动作组合。开启链式推理后,两市场的"
        "持有不动占比从不足 1% 上升到 5.6%–7.3%,信念更新占比从"
        "约 64% 降至约 53%–56%,撤单占比从 8.5%–8.8% 降至约 6.3%;"
        "限价挂单占比略升(22.2%–22.7% → 25.3%–26.7%)。也就是说,"
        "推理过程使更多决策落到“不动作”和“限价挂单”上,但这种"
        "动作结构变化并没有对应一致方向的价格变化。")
    three_line_table(doc, "显式推理实验核心指标",
        ["市场", "显式推理", "终态价", "成交额", "持有不动", "信念更新", "撤单"],
        [
            ["Robotaxi", "开启", "0.417 ± 0.033", "800", "7.3",  "52.8", "6.3"],
            ["Robotaxi", "关闭", "0.443 ± 0.020", "483", "0.9",  "63.7", "8.8"],
            ["Ethereum", "开启", "0.553 ± 0.016", "817", "5.6",  "55.9", "6.3"],
            ["Ethereum", "关闭", "0.568 ± 0.035", "786", "0.3",  "64.8", "8.5"],
        ],
        widths_cm=[2.0, 2.6, 2.6, 1.6, 1.8, 2.8, 1.8])
    para(doc,
        "价格指标再次提示本文保持谨慎。Robotaxi 上推理开启和关闭"
        "的终价分别为 0.417 与 0.443,Ethereum 上为 0.553 与 0.568,"
        "两市场差异(0.026、0.015)均落在随机性基线之内。与之相比,"
        "动作占比差异更清楚:持有不动从关闭组的 0.3%–0.9% 升至"
        "开启组的 5.6%–7.3%,信念更新从 63.7%–64.8% 降至"
        "52.8%–55.9%,撤单从 8.5%–8.8% 降至 6.3%。因此,显式推理"
        "主要改变动作结构与活跃度,而不是稳定改变终态价格。")
    para(doc,
        "综合来看,信念状态和初始判断生成方式对行为质量和价格"
        "偏移最关键;画像分布与显式推理开关主要影响动作结构、交易"
        "活跃度或局部波动,但在当前样本下不能稳定解释终态价格方向。")
    para(doc,
        "模块实验说明,部分设计能够改善行为连续性或改变动作结构,但这些"
        "改进仍不能被直接解释为更强的结算方向预测能力。在此基础上,"
        "本文将系统放入尚未结算的活跃市场,不再评价命中率,而是检验它"
        "能否在真实结算出现前组织价格路径、交易压力和分歧区间,作为一种"
        "情景分析工具。")

    h2(doc, "五、活跃市场预演:可呈现结算前分歧,不宜作方向预测")
    para(doc,
        "本节结论:按当前实验快照,三组种子下形成完整价格路径,但终态 YES 价格分散于 "
        "0.495–0.620,适合作为结算前情景证据,不能据此作出方向预测。"
        "活跃市场预演使用 Thunder NBA Finals 这一市场,设置 20 个智能体、"
        "20 轮和三个随机种子。由于市场尚未结算,本文只报告结算前情景"
        "路径,不评价预测命中率。三组种子均从约 0.575 起步,但终态"
        "方向并不一致:终价 0.495 / 0.620 / 0.555,均值约 0.557,"
        "标准差约 0.063。")
    figure(doc, "开放市场_结算前情景价格路径.png",
           "活跃市场三种子 YES 中间价路径", width_cm=12)
    para(doc,
        "三条结算前路径都从 0.575 起步,但终点分别落在 0.495、"
        "0.620 与 0.555,跨度达到 0.125。最大波动率出现在种子 2"
        "(0.0153),最小出现在种子 1(0.0073)。这说明活跃市场实验"
        "能够生成多条可读的情景路径,但路径方向并不一致:三条路径"
        "分别出现 −0.080、+0.045 与 −0.020 的漂移,与结束市场样本中"
        "方向不稳定的现象相互呼应。")
    three_line_table(doc, "活跃市场预演各种子指标",
        ["随机种子", "起始价", "终态价", "漂移", "波动率", "成交笔数"],
        [
            ["0", "0.575", "0.495", "−0.080", "0.0100", "67"],
            ["1", "0.575", "0.620", "+0.045", "0.0073", "74"],
            ["2", "0.575", "0.555", "−0.020", "0.0153", "69"],
        ],
        widths_cm=[1.8, 2.0, 2.0, 2.0, 2.0, 2.0])
    para(doc,
        "三组运行的成交笔数都在 67–74 之间,说明活跃市场预演并非"
        "停留在价格曲线层面,而是产生了足够的撮合记录。跨种子终价"
        "标准差约 0.063,大于本文随机性基线的 ±0.045,进一步表明"
        "该题目下智能体群体存在较大分歧。由于该市场尚未结算,这些"
        "数字不能与真实方向比对,只能用于描述结算前情景和分歧区间。")
    para(doc,
        "活跃市场实验表明,LLM 驱动智能体可以提供结算前"
        "情景证据、交易压力和分歧区间,但该结果应等待市场关闭后再"
        "评价准确性。该体育市场上智能体群体方向分歧较大,仿真结果"
        "宜理解为结算前情景与分歧的呈现,而非收敛的预测。")
    para(doc,
        "无论是结束市场还是活跃市场,宏观结果都显示出同一个特征:系统"
        "能够形成交易过程,但价格方向并不稳定。要解释这一现象,仅看终态"
        "价格和成交量还不够,必须进一步进入微观层面,观察智能体声明的"
        "信念是否真正影响交易方向,以及群体信念如何与市场价格发生偏离。")

    h2(doc, "六、决策链分析:个体按信念交易,群体信念相对市价偏悲观")
    para(doc,
        "本节结论:个体交易方向与自报信念高度一致,但群体平均信念系统性"
        "低于市场中间价,智能体更像按自身判断交易,而非简单跟随现价。"
        "为考察上述宏观结果背后的微观机制,本文对代表性"
        "仿真样本开展决策链分析,聚合 36 次仿真共 720 个智能体的"
        "信念-行为序列。两个发现尤其值得指出。第一,智能体内部高度"
        "自洽:平均 94.2% 的交易方向与该智能体最近一次声明的信念"
        "一致,36 次仿真的一致率均明显高于 50% 随机基线。"
        "第二,群体信念存在系统性偏离市场价格的现象:平均“信念 −"
        "价格”差为 −0.186,意味着智能体群体集体比市场更悲观。"
        "Ethereum 上这一偏离最强(达 −0.21 至 −0.39),正好对应"
        "研究问题三中观察到的“长时模拟下 Ethereum 价格朝 0 漂移”。")
    figure(doc, "决策链_信念与交易方向一致性.png",
           "信念与交易方向一致性及信念—价差分布",
           width_cm=15)
    para(doc,
        "信念与交易方向的高度一致,说明智能体不是随机下单。6 次基线"
        "仿真的一致率在 0.879 至 0.991 之间,均值 0.946,平均每次"
        "仿真 121 笔交易。进一步看,一致交易多出现在 |信念−价格| "
        "较大的区域,尤其是在智能体认为价格明显偏离自身判断时。"
        "也就是说,智能体不仅声明信念,还会把该信念转化为交易方向。")
    figure(doc, "决策链_智能体信念轨迹与市场价格.png",
           "两基底市场：智能体信念轨迹与市场价、群体均值信念",
           width_cm=15)
    para(doc,
        "群体信念与市场价之间的差异解释了前文价格下行的来源。"
        "Tesla Robotaxi 上群体均值信念为 0.419,与同期市场价"
        "0.40–0.45 区间接近;Ethereum 上群体均值信念为 0.474,"
        "低于同期市场价 0.52–0.56,差距约 0.05–0.09。单个智能体"
        "的信念轨迹整体平稳,中途偶有小幅修正,说明智能体并非每轮"
        "都剧烈改变看法,而是在多数轮次中维持原始判断。")
    figure(doc, "决策链_群体信念与市场价格差异时间演化.png",
           "群体信念与市场价的差异随轮次变化",
           width_cm=15)
    para(doc,
        "从时间序列看,这种偏离并非一次性出现。Robotaxi 上“群体"
        "均值信念 − 市场中间价”在 −0.13 至 −0.04 之间,均值为"
        "−0.040;Ethereum 上从初始的 +0.003 单调下行到末轮的"
        "−0.363,均值为 −0.090。两市场上差值整体为负,说明智能体"
        "群体比市场更悲观,且 Ethereum 的偏离随时间持续扩大。")
    figure(doc, "决策链_群体信念离散度随轮数变化.png",
           "群体信念离散度随轮次变化", width_cm=15)
    para(doc,
        "信念离散度下降说明群体存在弱趋同。Robotaxi 上初始信念"
        "离散度 0.144,末轮降至 0.086,均值 0.112;Ethereum 上"
        "初始 0.191,末轮 0.138,均值 0.179。两市场上离散度都随"
        "轮数下行,但 Ethereum 的离散度始终高于 Robotaxi,与该市场"
        "上更强的种子分歧一致。也就是说,智能体之间的看法会逐步"
        "收窄,但并未完全收敛到统一判断。")
    figure(doc, "决策链_智能体信念分布始末对比.png",
           "智能体在首次声明与最后声明时的信念分布", width_cm=15)
    para(doc,
        "始末分布进一步显示,智能体信念在多轮交易后整体被压低。"
        "Robotaxi 上首次声明均值为 0.405、末次为 0.362,下移约"
        "0.04;Ethereum 上首次均值为 0.568、末次为 0.202,下移约"
        "0.37。两市场上分布形态从首次的较宽分布(标准差 0.154/0.193)"
        "向末次的更窄分布(0.095/0.138)收敛。其中 Ethereum 的压低"
        "幅度尤其大,与该市场上长轮次价格下行同源。")
    figure(doc, "决策链_智能体信念更新次数分布.png",
           "每个智能体的信念更新次数分布", width_cm=15)
    para(doc,
        "信念更新次数说明显式信念机制被稳定使用。Robotaxi 上更新"
        "次数均值 12.4、中位数 13.0、范围 5–19;Ethereum 上均值"
        "12.9、中位数 12.5、范围 4–18。20 轮仿真下,多数智能体在"
        "大约 12–13 轮显式更新信念,即平均每两轮更新一次,极少有"
        "智能体几乎不更新或每轮都更新。")
    para(doc,
        "围绕动作结构,本文记录了每个智能体的动作占比、相邻轮次"
        "动作转移矩阵,以及全部规模实验中按规模分组的撤单率轨迹。"
        "智能体之间存在明显的“专精偏好”:同一仿真中,多数智能体以"
        "信念更新与限价挂单为主导动作,少数智能体倾向于"
        "高频份额拆分或市价成交操作。第一笔交易延迟集中在前 4–6 轮,"
        "之后逐步进入稳定的信念声明、挂单和撤单循环。")
    figure(doc, "决策链_智能体动作偏好热图.png",
           "智能体动作占比热图（按主导动作排序）", width_cm=15)
    para(doc,
        "动作结构方面,多数智能体以信念更新和限价挂单为主。整体"
        "平均上,信念更新占比 55.3%、限价挂单占 26.6%、持有不动占"
        "7.7%、撤单占 6.3%、份额拆分占 2.9%、市价成交占 1.2%,"
        "份额合并几乎为零。与此同时,少数智能体在份额拆分或市价"
        "成交上占比更高,形成一定的“专精偏好”。这表明智能体之间"
        "存在动作组合差异,并没有被语言模型拉成完全相同的行为模式。")
    figure(doc, "决策链_动作组成随轮数演化.png",
           "动作组成随轮次演化", width_cm=15)
    para(doc,
        "动作组成在仿真过程中并非均匀分布,而是随轮数缓慢变化。"
        "早期(第 0–3 轮)以信念更新与限价挂单为主,撤单占比较高;"
        "中后期份额拆分与市价成交出现稳定但小幅度的占比,持有不动"
        "在开启显式推理的仿真中保持约 5%–7%。这一变化说明前期"
        "更像试探挂单,中后期在信念相对稳定后进入持续的声明、挂单"
        "和撤单循环。")
    figure(doc, "决策链_相邻轮次动作转移矩阵.png",
           "相邻轮次动作转移概率热图", width_cm=12)
    para(doc,
        "相邻动作之间存在清楚的时序关系。对角线较深,说明智能体倾向"
        "在相邻轮次重复相同动作类别;信念更新与限价挂单之间存在"
        "明显的双向转移,与“先声明信念、再据此挂单”的链路一致。"
        "撤单之后下一步多回到信念更新或限价挂单,而不是再次撤单。"
        "因此,动作序列不是离散动作的简单堆积,而具有可解释的决策顺序。")
    figure(doc, "决策链_智能体首次交易轮数分布.png",
           "首次交易轮次分布", width_cm=15)
    para(doc,
        "首次成交时间集中在前几轮。Robotaxi 上首次成交轮次中位数"
        "2.0、均值 1.92、最大 7;Ethereum 上中位数 1.0、均值 1.37、"
        "最大 6。也就是说,在 20 轮仿真中,多数智能体在前 2 轮就"
        "完成第一笔成交,Ethereum 比 Robotaxi 更早进入交易状态。")
    figure(doc, "规模效应_单笔成交名义额分布.png",
           "单笔成交名义额分布（n=50）", width_cm=15)
    para(doc,
        "单笔成交规模呈现明显右偏。Robotaxi 上成交笔数 489,均值"
        "43.2 USD,中位数 10.0 USD,标准差 162;Ethereum 上成交笔数"
        "517,均值 68.8 USD,中位数 10.0 USD,标准差 308,最大值达"
        "5000 USD。多数成交规模很小,少数大额成交贡献了大部分名义额。"
        "这一分布特征与真实预测市场中“多小单、少量大单”的形态接近,"
        "也解释了为何损益分布尾部主要由少数大额成交主导。")
    figure(doc, "模块消融_不同行为画像的个体盈亏分布.png",
           "六类行为画像下的智能体盈亏分布",
           width_cm=15)
    para(doc,
        "画像分组后的个体盈亏进一步说明,画像对少数极端个体更有"
        "解释力,对多数智能体的区分度有限。Robotaxi 上 C0(广覆盖高活跃)"
        "与 C5(大额配置)均值"
        "为 +16.3 与 +23.8 USD,C4(高赔率追逐)为 −9.2 USD;Ethereum"
        "上 C2(高频套利)与 C5 均值最高(+84.8 与 +11.1 USD),C3"
        "(低价逆向)为 −13.2 USD。多数画像的盈亏中位数仍接近 0,"
        "差异主要体现在均值与极端值上,说明画像在解释“平均赢家/输家”"
        "时有一定区分度,但在解释多数智能体的盈亏时区分度有限。")
    figure(doc, "决策链_声明信念与置信度关系.png",
           "声明 YES 概率与置信度散点",
           width_cm=12)
    para(doc,
        "自报置信度没有随着仿真推进呈现单调变化。Robotaxi 显式推理"
        "开启基线中的 589 次信念声明显示,YES 概率跨度为 0.05 至"
        "0.99,置信度均值约 0.578;声明点在 YES 概率 0.35–0.60 区间"
        "最密集,置信度多数集中在 0.4–0.7。前期与后期声明在置信度"
        "上没有明显分层,说明智能体并未集体提高或降低自报置信度。"
        "这一结果为后续讨论“自报置信度能否作为交易权重”提供了依据。")
    para(doc,
        "综合上述微观分析,本文得到的一般性结论是:大语言模型驱动智能体"
        "在交易场景中具备较强的“言行一致性”——它的交易方向几乎总是"
        "与它自己声明的信念吻合;但群体层面会出现系统性偏离市场价格"
        "的现象,这一偏离与具体市场和真值方向有关。从群体层面看,智能体"
        "不是“跟随价格”的执行者,而是“按自己的看法交易”的主体,"
        "这一结果对行为模拟研究具有解释价值。")

    h2(doc, "七、本章小结")
    para(doc,
        "本章围绕结束市场复现、规模实验、轮数实验、模块消融、活跃市场"
        "预演和决策链分析,检验了大语言模型驱动智能体在 Polymarket "
        "订单簿环境中的行为模拟能力。总体来看,智能体能够在资金、持仓、"
        "价格刻度和撮合规则约束下形成完整交易过程,动作、成交、持仓、"
        "价格和信念记录均可追踪,说明系统具有可执行和可审计的行为模拟"
        "能力。但是,这一能力不能被扩大解释为稳定的结算方向预测能力。"
        "结束市场中,终态价格位于胜方一侧与模拟过程中朝胜方方向移动"
        "并不是同一件事,后者更能反映过程性价格发现,而本文实验显示"
        "这种价格发现能力仍然有限。")
    para(doc,
        "进一步的受控实验表明,智能体数量和模拟轮数主要改变交易强度、"
        "反馈长度和路径依赖,而不是稳定重塑价格方向。规模扩大显著提高"
        "成交笔数和成交名义额,并放大个体盈亏尾部;轮数增加则累积更多"
        "交易机会,但在部分市场上也放大种子差异和路径漂移。模块实验"
        "显示,显式信念机制与初始主观判断生成方式对行为质量影响更清楚,"
        "画像分布和显式推理开关则更多改变动作结构、活跃度或局部波动,"
        "当前样本不足以支持其稳定改善价格发现的结论。")
    para(doc,
        "活跃市场预演进一步说明,该系统适合用于结算前情景分析,而不宜"
        "直接用于方向预测。三组种子能够生成完整的价格路径和成交记录,"
        "但终态分布存在明显分歧,只能解释为价格区间、交易压力和群体"
        "分歧的展示。决策链分析则从微观层面解释了这一边界:智能体个体"
        "交易方向与自报信念高度一致,说明其行为并非随机;但群体信念相对"
        "市场价格存在系统性偏离,因此个体层面的自洽不会自动转化为群体"
        "层面的正确价格发现。由此,本章将本文模拟系统的适用范围界定为"
        "预测市场行为模拟、机制对照和微观结构分析,而不是替代真实市场"
        "作出结算判断。")

    # ====================================================================
    # 第五章 总结与讨论
    # ====================================================================
    pagebreak(doc)
    h1(doc, "第五章  总结与讨论")
    para(doc,
        "本章在前文实证分析的基础上,归纳主要研究结论,说明研究局限,"
        "并提出未来研究方向。")

    h2(doc, "一、主要研究结论")
    para(doc,
        "本文围绕“构建可用于去中心化预测市场的智能体交易模拟系统”这一"
        "目标,在 Polymarket 真实数据校准的环境中,检验大语言模型驱动"
        "智能体能否承担可执行交易者角色,并分析智能体规模、模拟轮数及"
        "关键模块设计对市场结果的影响。本文不是把智能体输出作为一次性"
        "预测答案,而是把下单、成交、撤单、持仓、信念更新和价格路径作为"
        "共同证据,观察个体行动如何在订单簿制度下汇聚为市场过程。")
    para(doc,
        "围绕研究问题一,10 个结束市场、30 组仿真均形成完整的动作、成交、"
        "持仓与价格记录;按本地结算侧标注,21 组终态位于胜方一侧,但仅 13 组"
        "在过程中朝胜方方向移动,说明终态是否“看起来正确”与过程是否发生"
        "价格发现不能混为一谈。决策链分析显示,个体交易方向与自报信念高度"
        "一致,而群体信念相对市场价格存在系统性偏低。")
    para(doc,
        "围绕研究问题二与三,规模与轮数实验表明,智能体数量与模拟轮数主要"
        "作用于交易强度与反馈长度,而非稳定重塑价格方向。围绕研究问题四,"
        "显式信念机制与初始主观判断生成方式对行为质量的影响相对明确;"
        "画像分布对终态方向的影响不稳定,显式推理开关更多体现在动作结构"
        "与成交笔数上。围绕研究问题五,Thunder NBA Finals 活跃市场预演在"
        "三组种子下终态价格分别为 0.495、0.620 与 0.555,跨种子方向并不一致,"
        "仿真更适合呈现结算前分歧而非给出单一预测。")
    para(doc,
        "综上,本文构建的模拟框架能够支持去中心化预测市场的行为模拟与"
        "机制对照,其合理用途是观察市场微观过程与设计要素的影响,而不是"
        "替代真实市场作出结算判断。")

    h2(doc, "二、研究局限")
    para(doc,
        "首先,实验重复次数有限。受语言模型调用成本与运行时间制约,多数"
        "受控实验仅设置三个随机种子;对画像分布、长时 Ethereum 漂移及活跃"
        "市场方向等结果仍只能作初步判断。其次,市场样本范围仍偏窄,结束"
        "市场样本虽覆盖多种题材,但相对 Polymarket 全量市场而言规模有限,"
        "受控实验亦集中在 Robotaxi 与 Ethereum 两个基底市场。再次,模拟环境"
        "对真实机制作了必要简化,外部交易者进入、实时新闻、跨市场套利及"
        "链上延迟等因素尚未纳入;画像、信念与理由来自模型输出与实验构造,"
        "不能等同于真实交易者心理。")

    h2(doc, "三、未来研究方向")
    para(doc,
        "在数据与样本方面,可扩大结束市场样本,并在更多基底市场上重复规模、"
        "轮数与模块实验;同时可引入链上日志、平台评论或外部新闻,使信息环境"
        "更接近真实预测市场。在模型与实验设计方面,可进一步检验信息冲击"
        "强度、信念更新频率、参与者进入与退出等机制,并结合更大样本与更多"
        "随机重复提高统计功效。在应用与工具化方面,可开发面向研究者的实验"
        "界面,分别服务于机制研究、教学演示与活跃市场的结算前情景分析;"
        "评价指标亦可从价格与成交扩展至价差恢复、流动性消耗、信念校准误差等。")
    # ---------- 参考文献 ----------
    pagebreak(doc)
    h1(doc, "参考文献")
    for r in REFERENCES:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.left_indent = IND_HANG          # 悬挂缩进
        pf.first_line_indent = -IND_HANG
        run = p.add_run(r)
        _set_run(run, ea=SONG, size=PT_BODY)

    # ---------- 致谢 ----------
    pagebreak(doc)
    h1(doc, "致  谢")
    para(doc,
        "感谢导师在选题、方法与写作上的悉心指导;感谢提供公开数据与"
        "文档的开源社区;感谢同窗的讨论与建议。文中所有结论与不足"
        "由作者负责。")

    # ---------- 附录 ----------
    pagebreak(doc)
    h1(doc, "附录 A  市场用户行为画像的数据定义")
    para(doc,
        "本附录说明钱包行为画像聚类的关键参数与运行结果。每个市场的"
        "事件前钱包池首先经过 8 个行为特征的标准化处理:累计名义"
        "交易额、最大单一市场集中度、单位资金市场广度、平均成交价、"
        "尾部极端价格交易占比、活跃时长、成交价波动,以及 burstiness"
        "(交易时间集中程度)。所有特征再被 winsorize 至 [p1, p99] "
        "以抑制重尾。聚类采用 K-means,扫描 K∈{3,4,5,6},每个 K 用"
        "50 次自助重抽样估计 Jaccard 稳定性,并选取满足“轮廓系数 ≥"
        "0.20、Jaccard 稳定性 ≥ 0.75、最小簇占比 ≥ 3%”的最大合格"
        "K。在本文用于画像构建的十三个市场上,该方案稳定选出 K=6,最小"
        "簇占比 3.4%–5.2%,自助 Jaccard 0.77–0.99。")
    para(doc,
        "需特别说明的是,该聚类是数据驱动的统计概括,而非研究者"
        "手工设定的角色。每一类画像由一组在 8 个行为特征上相互接近"
        "的真实钱包构成,其特征中心即为该画像的定量定义。正文 §4 中"
        "画像分布实验进一步证实,改变画像在群体中的占比对市场层面结果的"
        "影响在当前样本下不稳定——也就是说,画像是对参与者总体的"
        "描述性分类,而不是决定市场结果的因果机制。")

    h1(doc, "附录 B  实验复现说明")
    para(doc,
        "本文所有量化结论均可由随附的实验配置与分析制品复现:每一仿真"
        "完整记录其动作序列、成交、持仓与盘口轨迹;受控对照的设计、"
        "随机种子与导出参数均被固定记录。决策温度为零,语言模型后端"
        "为唯一的外部随机来源,其残余非确定性已在随机性基线中量化"
        "(终价 ±0.045),本文所有显著性判定均在该噪声带之上做出。")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    default_out = (Path.home() / "Desktop" /
                   "大语言模型驱动智能体的行为模拟研究——以Polymarket预测市场为例.docx")
    ap.add_argument("--out", default=str(default_out))
    ap.add_argument(
        "--template",
        default="/Users/moonshot/Projects/Poly/参考/附件1：2026届本科毕业论文工作手册/"
                "17-南京大学本科毕业论文模板【WORD版本】2026届更新.docx",
        help="NJU template .docx; falls back to a blank document if absent",
    )
    args = ap.parse_args()

    tpl = Path(args.template)
    if tpl.exists():
        doc = Document(str(tpl))
        # clear the template's placeholder body, keep its styles
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = Document()

    _page_setup(doc)
    build(doc)
    out = Path(args.out).expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"wrote {out}  ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
