#!/usr/bin/env python3
"""
build_thesis_docx.py
读取 docs/polished/ 下的润色 Markdown 文件，生成符合南京大学格式的 Word 论文。

核心功能：
  - Markdown 表格 → 三线表（自动去掉分隔行 | --- |）
  - 图注占位行（以"图N-M"开头的独立行）→ 插入 docs/figures/figN_M.png
  - 粗体 **...** → 黑体/加粗行内文字
  - ## / ### 标题 → Word 样式标题

用法：
  python build_thesis_docx.py [--out 输出路径.docx]
"""

import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 路径 ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent
POLY_DIR     = SCRIPT_DIR.parent                        # polymetl/
POLISHED_DIR = POLY_DIR / "docs" / "polished"
FIG_NEW      = POLY_DIR / "docs" / "figures"            # fig4_1.png 等
FIG_OLD      = POLY_DIR / "docs" / "v14" / "figures"   # fig1_loop.png 等

CHAPTER_FILES = [
    "ch1_polished.md",
    "ch2_polished.md",
    "ch3_polished.md",
    "ch4a_polished.md",
    "ch4b_polished.md",
    "ch5_polished.md",
]

# ── 字体 / 字号常量 ────────────────────────────────────────────────────────────
SONG  = "宋体"
HEI   = "黑体"
KAI   = "楷体_GB2312"
LATIN = "Times New Roman"
PT_BODY  = 12    # 小四
PT_H1    = 16    # 三号
PT_H2    = 14    # 四号
PT_H3    = 12    # 小四黑体
PT_CAP   = 10.5  # 图/表题注
IND_BODY = Cm(0.85)   # 首行缩进

# 全局计数器
_fig_counter = {"n": 0}
_tbl_counter = {"n": 0}

# ── 底层辅助 ──────────────────────────────────────────────────────────────────
def _set_run(run, *, ea=SONG, latin=LATIN, size=PT_BODY, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"),   latin)
    rf.set(qn("w:hAnsi"),   latin)
    rf.set(qn("w:eastAsia"), ea)


def _set_cell_border(cell, **edges):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, props in edges.items():
        if props is None:
            continue
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"),  props.get("val", "single"))
        el.set(qn("w:sz"),   str(props.get("sz", 6)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


# ── 段落添加函数 ──────────────────────────────────────────────────────────────
def add_heading(doc, text: str, level: int):
    """level 1=章，2=节，3=小节"""
    p    = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf   = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(0)

    size = {1: PT_H1, 2: PT_H2, 3: PT_H3}.get(level, PT_H3)
    run  = p.add_run(text)
    _set_run(run, ea=HEI, size=size, bold=True)
    return p


def add_body_para(doc, inline_parts, *, indent=True, center=False):
    """
    inline_parts: list of (text, bold) tuples
    bold=True → 黑体加粗；False → 宋体正文
    """
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.line_spacing     = Pt(22)
    pf.space_before     = Pt(0)
    pf.space_after      = Pt(0)
    pf.first_line_indent = IND_BODY if indent else Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY

    for text, bold in inline_parts:
        run = p.add_run(text)
        _set_run(run, ea=HEI if bold else SONG, bold=bold)
    return p


def add_three_line_table(doc, title: str, headers: list, rows: list,
                          widths_cm: list = None):
    """三线表：上粗线 → 表头 → 中细线 → 数据行 → 下粗线"""
    _tbl_counter["n"] += 1
    # 表题（上方居中）
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr  = cap.add_run(f"表{_tbl_counter['n']}  {title}")
    _set_run(cr, ea=SONG, size=PT_CAP, bold=True)

    thick = {"val": "single", "sz": 12}
    thin  = {"val": "single", "sz": 6}
    n_cols = len(headers)
    n_data = len(rows)

    t = doc.add_table(rows=1 + n_data, cols=n_cols)
    t.alignment = 1  # center

    # 表头行
    for ci, htext in enumerate(headers):
        cell = t.rows[0].cells[ci]
        cell.text = str(htext)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                _set_run(run, ea=SONG, size=10, bold=True)
        _set_cell_border(cell, top=thick, bottom=thin)

    # 数据行
    for ri, row in enumerate(rows, start=1):
        is_last = (ri == n_data)
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_run(run, ea=SONG, size=10)
            if is_last:
                _set_cell_border(cell, bottom=thick)

    # 列宽
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm[:n_cols]):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    return t


def add_figure(doc, png_path: Path, caption: str, width_cm: float = 13.5):
    """插入图片，题注置于图下方"""
    _fig_counter["n"] += 1
    if not png_path.exists():
        # 占位段落，避免因缺图崩溃
        p = doc.add_paragraph(f"[图片未找到：{png_path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(png_path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr  = cap.add_run(caption)
    _set_run(cr, ea=SONG, size=PT_CAP, bold=True)
    doc.add_paragraph()


# ── 行内富文本解析（**bold** 与普通文字混合）────────────────────────────────
def parse_inline(text: str):
    """
    返回 list of (text_segment, is_bold) tuples。
    支持 **...**  形式的粗体标记。
    """
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts if parts else [(text, False)]


# ── 图注行检测与 PNG 路径解析 ────────────────────────────────────────────────
# 匹配形如 "图4-1  ..." 或 "图 4-1  ..." 开头的独立行
_FIG_LINE_RE = re.compile(r'^图\s*(\d+)-(\d+)\s+(.+)$')


def resolve_figure_path(ch: str, num: str) -> Path:
    """
    优先在 docs/figures/ 中查找 figCH_NUM.png；
    若不存在，回退到 docs/v14/figures/。
    """
    name_new = f"fig{ch}_{num}.png"
    p = FIG_NEW / name_new
    if p.exists():
        return p
    # 回退：扫描 v14/figures 目录找近似名
    if FIG_OLD.exists():
        for f in FIG_OLD.iterdir():
            if f.suffix.lower() == ".png" and f.stem.startswith(f"fig{ch}"):
                return f
    return FIG_NEW / name_new   # 不存在时仍返回，add_figure 会显示占位


# ── Markdown 解析器 ───────────────────────────────────────────────────────────
def parse_markdown(md_text: str):
    """
    将 Markdown 文本解析为 block 列表。
    每个 block 为 dict，type 字段为：
      heading   → {level, text}
      figure    → {ch, num, caption}
      table     → {title, headers, rows}
      paragraph → {parts: [(text, bold)]}
      rule      → {}        （--- 分隔线）
      blank     → {}
    """
    lines  = md_text.splitlines()
    blocks = []
    i      = 0
    pending_table_title = None   # **表 X-X ...** 行暂存

    while i < len(lines):
        line = lines[i]

        # 空行
        if line.strip() == "":
            blocks.append({"type": "blank"})
            i += 1
            continue

        # 水平线
        if re.match(r'^-{3,}$', line.strip()):
            blocks.append({"type": "rule"})
            i += 1
            continue

        # 标题行（## / ###）
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            i += 1
            continue

        # 图注行（独立行，以"图N-M"开头）
        m = _FIG_LINE_RE.match(line.strip())
        if m:
            blocks.append({
                "type":    "figure",
                "ch":      m.group(1),
                "num":     m.group(2),
                "caption": m.group(3).strip(),
            })
            i += 1
            continue

        # 表格块（| ... | 开头）
        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            # 去掉分隔行（| --- | 形式）
            header_row = None
            data_rows  = []
            for tl in tbl_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                if all(re.match(r'^-+$', c) for c in cells if c):
                    continue   # 跳过分隔行
                if header_row is None:
                    header_row = cells
                else:
                    data_rows.append(cells)
            if header_row:
                title = pending_table_title or ""
                pending_table_title = None
                blocks.append({
                    "type":    "table",
                    "title":   title,
                    "headers": header_row,
                    "rows":    data_rows,
                })
            continue

        # 粗体独立行作为表格标题暂存（**表 X-X ...**）
        m_tbl = re.match(r'^\*\*表\s*(.+?)\*\*\s*$', line.strip())
        if m_tbl:
            pending_table_title = m_tbl.group(1).strip()
            i += 1
            continue

        # 普通段落（含行内粗体）
        blocks.append({"type": "paragraph", "parts": parse_inline(line.strip())})
        i += 1

    return blocks


# ── 渲染器 ────────────────────────────────────────────────────────────────────
# 需要特殊处理的节标题关键词
_SKIP_TOC      = re.compile(r'^目录')
_ABSTRACT_CN   = re.compile(r'^摘要')
_ABSTRACT_EN   = re.compile(r'^Abstract')
_SPECIAL_H1    = re.compile(r'^(致谢|附录)')
_TITLE_LINE    = re.compile(r'^#\s')   # 文件级 # 标题，作为文档大标题跳过


def render_blocks(doc, blocks):
    """将 block 列表渲染到 doc。"""
    skip_toc_body = False   # 跳过目录正文

    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue   # 不输出多余空行

        elif btype == "rule":
            doc.add_paragraph()   # 用空段落模拟间距
            continue

        elif btype == "heading":
            level = block["level"]
            text  = block["text"]

            # 文档级 # 标题：跳过（由封面处理）
            if level == 1 and not text.startswith("第"):
                # 摘要、Abstract、目录等
                if _SKIP_TOC.match(text):
                    skip_toc_body = True
                    continue
                elif _ABSTRACT_CN.match(text) or _ABSTRACT_EN.match(text):
                    skip_toc_body = False
                    add_heading(doc, text, 1)
                elif _SPECIAL_H1.match(text):
                    skip_toc_body = False
                    add_heading(doc, text, 1)
                else:
                    skip_toc_body = False
                    add_heading(doc, text, 1)
                continue

            skip_toc_body = False
            # 章节标题对应 level
            doc_level = min(level, 3)
            add_heading(doc, text, doc_level)

        elif btype == "figure":
            if skip_toc_body:
                continue
            png = resolve_figure_path(block["ch"], block["num"])
            caption = f"图{block['ch']}-{block['num']}  {block['caption']}"
            add_figure(doc, png, caption)

        elif btype == "table":
            if skip_toc_body:
                continue
            add_three_line_table(
                doc,
                title   = block["title"],
                headers = block["headers"],
                rows    = block["rows"],
            )

        elif btype == "paragraph":
            if skip_toc_body:
                continue
            parts = block["parts"]
            # 跳过目录正文（形如"　　一、研究背景"的缩进行）
            raw = "".join(t for t, _ in parts)
            if not raw.strip():
                continue
            add_body_para(doc, parts)


# ── 主文档构建 ────────────────────────────────────────────────────────────────
def build(doc):
    """读取各章 Markdown 并逐章渲染。"""
    for fname in CHAPTER_FILES:
        fpath = POLISHED_DIR / fname
        if not fpath.exists():
            print(f"  [警告] 文件不存在，跳过：{fpath}", file=sys.stderr)
            continue

        text   = fpath.read_text(encoding="utf-8")
        blocks = parse_markdown(text)
        render_blocks(doc, blocks)

        # 章末分页（每个 ch 文件之间插入分页符）
        doc.add_page_break()


def main():
    ap = argparse.ArgumentParser(description="将润色 Markdown 章节渲染为 Word 论文")
    ap.add_argument(
        "--out", default=None,
        help="输出 .docx 路径（默认：Poly/ 目录下）"
    )
    args = ap.parse_args()

    out_path = Path(args.out) if args.out else (
        POLY_DIR.parent / "大语言模型驱动智能体的去中心化预测市场行为模拟研究.docx"
    )

    print("正在构建论文 Word 文档……")
    doc = Document()

    # 全局页面设置（A4，页边距 2.5/3.0 cm）
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # 默认段落样式（正文）
    style = doc.styles["Normal"]
    style.font.name = LATIN
    style.font.size = Pt(PT_BODY)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    build(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已保存至：{out_path}")
    print(f"  图片计数：{_fig_counter['n']}  表格计数：{_tbl_counter['n']}")


if __name__ == "__main__":
    main()
